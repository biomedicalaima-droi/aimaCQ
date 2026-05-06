import streamlit as st
import numpy as np
import io
import os
from datetime import date
from typing import Dict, Any, List

import streamlit as st
from fpdf import FPDF
from datetime import date
import os
import tempfile
import time
import io
import sys
import base64
from PIL import Image
import pdfplumber

# Importations pour DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT 

# Importations pour PDF (Reportlab)
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader # NOUVEAU


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)
AIMA_LOGO_PATH = resource_path("aima_logo.png")
#AIMA_LOGO_PATH = "C:/Users/perso/Desktop/aima_logo.png"

AIMA_LOGO_PATH = "aima_logo.png" 
BOTTOM_LOGOS_PATH = "BOTOM_LOGOS.png"

# --- INITIALISATION SESSION STATE ---
if 'manual_items_dict' not in st.session_state:
    st.session_state.manual_items_dict = []
if 'active_catalog' not in st.session_state:
    st.session_state.active_catalog = []
if 'catalog_selector' not in st.session_state:
    st.session_state.catalog_selector = []
# On sépare les compteurs par type de document pour chaque client
if 'counters' not in st.session_state:
    st.session_state.counters = {"DEVIS": {}, "FACTURE": {}}

# --- STYLE CSS ---
st.markdown("""
    <style>
    .block-container { padding-top: 5rem; padding-bottom: 3rem; }
    h1 { margin-top: -10px !important; padding-bottom: 10px !important; }
    h3 { margin-top: -10px !important; margin-bottom: 10px !important; }
    hr { margin-top: 0px !important; margin-bottom: 10px !important; }
    </style>
    """, unsafe_allow_html=True)

# --- FONCTIONS UTILES ---
def get_base64_logo(path):
    if os.path.exists(path):
        try:
            with open(path, "rb") as f:
                data = base64.b64encode(f.read()).decode("utf-8")
            return f"data:image/png;base64,{data}"
        except Exception: return None
    return None

def import_items_from_pdf(uploaded_pdf):
    try:
        new_items = []
        with pdfplumber.open(uploaded_pdf) as pdf:
            for page in pdf.pages:
                table = page.extract_table()
                if table:
                    for row in table:
                        if not row or "Designation" in str(row[0]) or "TOTAL" in str(row[0]): continue
                        try:
                            nom = str(row[0]).strip()
                            prix_str = str(row[1]).replace(' ', '').replace('€', '').replace(',', '.')
                            prix = float(prix_str)
                            new_items.append({"id": str(time.time())+nom, "nom": nom, "prix": prix})
                        except: continue
        return new_items
    except Exception as e:
        st.error(f"Erreur d'importation : {e}")
        return []

# --- CALLBACKS ---
def delete_catalog_item(item_name):
    if item_name in st.session_state.catalog_selector:
        st.session_state.catalog_selector.remove(item_name)
    st.session_state.active_catalog = [x for x in st.session_state.active_catalog if x['name'] != item_name]

def delete_manual_item(index):
    st.session_state.manual_items_dict.pop(index)

# --------------------------------------------------------------------------------------
# --- NOUVEAU: CHEMINS DES LOGOS (À REMPLACER) ------------------------------------------
# --------------------------------------------------------------------------------------
# ATTENTION: Vous devez remplacer ces chemins par les chemins d'accès réels de vos images 
# lorsque vous exécutez le code sur votre machine locale.
# Pour Streamlit, ces fichiers doivent être accessibles.


# Logo AIMA (haut gauche)
AIMA_LOGO_PATH = "C:/Users/perso/Desktop/aima_logo.png" # REMPLACEZ PAR LE CHEMIN DU LOGO AIMA 

# Bloc de logos partenaires/certifications (bas)
BOTTOM_LOGOS_PATH = "C:/Users/TELEA133007WSM1/Desktop/BOTOM_LOGOS.png" # REMPLACEZ PAR LE CHEMIN DU BLOC DE LOGOS

ALTERNATIVE_NORMS = [
    "IEC 61010-1 (Exigences de sécurité pour les équipements électriques de laboratoire)",
    "IEC 62353 (Exigences de sécurité pour les équipements électriques de laboratoire)",
    "Autre norme de sécurité spécifique",
    "Non applicable",
    "Non applicable (Alimentation par batterie interne ou externe non-médicale)"
]

# --------------------------------------------------------------------------------------
# --- CONFIGURATION GLOBALE DES APPAREILS (10 DISPOSITIFS CLÉS) ------------------------
# --------------------------------------------------------------------------------------
# (Reste inchangé)

GLIDESCOPE_MONITOR = {
    "NAME": "Moniteur Vidéo Laryngoscope (GLIDESCOPE)",
    "TSE_REQUIRED": True,
    "FABRICANT_LIST": ["Verathon", "Autre"],
    "PERFORMANCE_CHECKS": {
        # La performance principale est qualitative (image), mais on peut tester l'autonomie
        "Autonomie Batterie": {"injected": 120, "tolerance": 15, "unit": "min", "type": "range"},
    },
    "SPECIFIC_CHECKS": {
        "Qualité d'Image (Moniteur)": "Image nette, couleur fidèle, sans pixels morts.",
        "Fonctionnement de l'Éclairage (Baton)": "Source lumineuse fonctionnelle et homogène.",
        "Test de Mise en Marche / Arrêt": "Mise en marche et arrêt du moniteur fonctionnels.",
        "Test du Bouton/Fonction d'Enregistrement": "Fonction d'enregistrement vidéo/image (si présente) opérationnelle.",
        "Compatibilité de la Lame/Baton": "Verrouillage et déverrouillage de la lame/du baton fonctionnel et stable.",
        "Indicateur de Charge Batterie": "Indicateur de niveau de batterie précis et fonctionnel.",
    },
    "VISUAL_CHECKS": [
        "Vérifier la propreté de l'appareil (boîtier, écran, connecteurs, socle de charge)",
        "Vérifier le bon état du boîtier du moniteur (pas de fissures, d'impacts)",
        "Vérifier l'état et l'intégrité du câble d'alimentation secteur.",
        "Vérifier le bon état et l'intégrité du Video Baton/câble réutilisable (absence de coupures ou de plis prononcés).",
        "Vérifier l'état de la fenêtre de la caméra (lentille) et de la source lumineuse (pas de rayures profondes ou d'opacités).",
        "Vérifier la lisibilité de l'étiquette constructeur (symboles Classe II et Type BF présents).",
    ],
    "SECURITY_CHECKS": {
        # Pas de test de Résistance de terre (PE) pour le CLASSE II (Double Isolation)
        # Courant de fuite Enveloppe (Châssis) - CLASSE II
        "Courant de fuite Enveloppe (NC)": {"limit": "< 500 μΑ (0.500 mA)", "unit": "mA", "type": "leakage"}, 
        # Courant de fuite Patient (NC) - TYPE BF
        "Courant de fuite Patient (NC)": {"limit": "< 100 μΑ (0.100 mA)", "unit": "mA", "type": "leakage"},
    }
}

# ---------------------------------------------------------------------------------------------------
# --- Dictionnaire Global QC_CONFIGS (POINT DE CORRECTION) -------------------------------------------
# ---------------------------------------------------------------------------------------------------

# La variable ECG_CONFIG_FULL n'est pas définie dans le bloc de code fourni, 
# nous supposons qu'elle est l'équivalent de l'entrée 'ECG'
# Pour que cela fonctionne, vous devez vous assurer que cette configuration existe.

ECG_CONFIG_FULL = {
    "NAME": "Électrocardiographe (ECG)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "GE Healthcare",
        "Philips Healthcare",
        "Schiller",
        "Mortara Instrument",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures quantitatives de précision)
    "PERFORMANCE_CHECKS": {
        
        # Vitesse de Défilement du Papier (Tolérances absolues)
        "Vitesse de défilement (15 mm/s)": {
            "injected": 15, 
            "tolerance": 0.5, 
            "unit": "mm/s", 
            "type": "absolute", 
            "consigne": "15 mm/s ± 0,5 mm/s"
        },
        "Vitesse de défilement (25 mm/s)": {
            "injected": 25, 
            "tolerance": 0.5, 
            "unit": "mm/s", 
            "type": "absolute", 
            "consigne": "25 mm/s ± 0,5 mm/s"
        },
        "Vitesse de défilement (50 mm/s)": {
            "injected": 50, 
            "tolerance": 1, 
            "unit": "mm/s", 
            "type": "absolute", 
            "consigne": "50 mm/s ± 1 mm/s"
        },
        
        # Amplitude du Signal (Sensibilité) (Tolérances absolues)
        "Amplitude du signal (5 mm/mV)": {
            "injected": 5, 
            "tolerance": 0.5, 
            "unit": "mm/mV", 
            "type": "absolute", 
            "consigne": "5 mm/mV ± 0,5 mm/mV"
        },
        "Amplitude du signal (10 mm/mV)": {
            "injected": 10, 
            "tolerance": 0.5, 
            "unit": "mm/mV", 
            "type": "absolute", 
            "consigne": "10 mm/mV ± 0,5 mm/mV"
        },
        "Amplitude du signal (20 mm/mV)": {
            "injected": 20, 
            "tolerance": 1, 
            "unit": "mm/mV", 
            "type": "absolute", 
            "consigne": "20 mm/mV ± 1 mm/mV"
        },
        
        # Fréquence Cardiaque (bpm) (Tolérances absolues)
        "Valeur de pouls (30 bpm)": {
            "injected": 30, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "30 bpm ± 5 bpm"
        },
        "Valeur de pouls (120 bpm)": {
            "injected": 120, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "120 bpm ± 5 bpm"
        },
        "Valeur de pouls (240 bpm)": {
            "injected": 240, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "240 bpm ± 5 bpm"
        },
        "Valeur de pouls (Néonat 94 bpm)": {
            "injected": 94, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "94 bpm ± 5 bpm"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Qualité du Tracé)
    "SPECIFIC_CHECKS": {
        
        "Autotest de l'appareil au démarrage": "Fonctionnel (Bon/Mauvais/NA)",
        "Fonctionnement des boutons (Marche/Arrêt, paramètres, alarme)": "Fonctionnel (Bon/Mauvais/NA)",
        "Passage automatique en mode batterie et fonctionnalité de la batterie": "Fonctionnel (Bon/Mauvais/NA)",
        "Vérification de l'enregistrement et de l'impression des paramètres": "Bon fonctionnement (Bon/Mauvais/NA)",
        "Représentation de la courbe normale (PQRS)": "Tracé correct (Bon/Mauvais)",
        "Représentation de la courbe anormale (PVC, etc.)": "Tracé correct (Bon/Mauvais)",
        "Vérification de la calibration du signal en mV (+/- 5%)": "Calibration correcte (Bon/Mauvais)",
        "Fréquence du signal et qualité du tracé (absence de bruit)": "Qualité optimale (Bon/Mauvais/NA)",
        "Vérification automatique des mesures (Calculs internes du dispositif)": "Calculs cohérents (Bon/Mauvais/NA)",
        "Contrôle des alarmes (Coupure alim, Électrode débranchée)": "Alarmes fonctionnelles (Bon/Mauvais/NA)",
    },
    
    # 3. Contrôles Visuels (État Général)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil",
        "Présence de tous les accessoires (câble secteur, câbles/capteurs patient, batterie, papier)",
        "Vérification de la lisibilité des affichages, voyants et claviers",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Basse Fréquence)
    "SECURITY_CHECKS": {
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # L'ECG est une partie appliquée de Type BF (Body Floating)
        "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
}


QC_CONFIGS = {
    "Électrocardiographe (ECG)": ECG_CONFIG_FULL,
    # --- AJOUT DU NOUVEAU DISPOSITIF ICI ---
    "GLIDESCOPE_MONITOR": GLIDESCOPE_MONITOR, 
    # --------------------------------------
    "Moniteur Patient": {
        "NAME": "Moniteur Patient Multiparamètres",
        "TSE_REQUIRED": True,
        "FABRICANT_LIST": ["Philips", "GE Healthcare", "Mindray", "Autre"],
        "PERFORMANCE_CHECKS": {
            "SpO2 90%": {"injected": 90, "tolerance": 2, "unit": "%", "type": "range"},
            "PNI 150 mmHg": {"injected": 150, "tolerance": 5, "unit": "mmHg", "type": "range"},
        },
        "SPECIFIC_CHECKS": {
            "Température (37.0°C)": "Affichage dans la tolérance (±0.2°C).",
            "Test d'Alarmes": "Alarmes audio/visuelles fonctionnelles.",
        },
        "VISUAL_CHECKS": ["Vérifier l’état des capteurs SpO2 et PNI", "Test de l'écran tactile"],
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"},
            "Courant de fuite patient": {"limit": "< 100 μΑ (0.100 mA)", "unit": "mA", "type": "leakage"},
            "Courant de fuite châssis": {"limit": "< 500 μΑ (0.500 mA)", "unit": "mA", "type": "leakage"},
        }
    },
    "Tensiomètre": {
        "NAME": "Tensiomètre Automatique (PNI)",
        "TSE_REQUIRED": True,
        "FABRICANT_LIST": ["Omron", "BPL", "Welch Allyn", "Autre"],
        "PERFORMANCE_CHECKS": {
            "Pression 50 mmHg": {"injected": 50, "tolerance": 3, "unit": "mmHg", "type": "range"},
            "Pression 200 mmHg": {"injected": 200, "tolerance": 3, "unit": "mmHg", "type": "range"},
        },
        "SPECIFIC_CHECKS": {
            "Test de Fuite du brassard": "Pression stable sur 30 secondes.",
            "Décompression Rapide (Dump)": "Décompression en moins de 10 secondes.",
        },
        "VISUAL_CHECKS": ["État du brassard et des tubes pneumatiques", "Vérifier la lisibilité de l'écran"],
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.200 Ω", "unit": "Ω", "type": "resistance"},
            "Courant de fuite châssis": {"limit": "< 500 μА (0.500 mA)", "unit": "mA", "type": "leakage"},
        }
    },
    
    
    "Inhalateur / Nébuliseur": {
        
    # Informations Générales
    "NAME": "Nébuliseur - Humidificateur",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Philips Respironics",
        "DeVilbiss Healthcare",
        "Fisher & Paykel Healthcare",
        "Omron",
        "Beurer",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Flux, Temps, Température)
    "PERFORMANCE_CHECKS": {
        
        # Température (Si équipé d'un réchauffeur/humidificateur. Tolérance standard: ± 2 °C)
        "Précision Température Réchauffeur (37 °C)": {
            "injected": 37.0, 
            "tolerance": 2.0, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "37.0 °C ± 2.0 °C après 5 min de chauffe"
        },
        
        # Flux d'air (Ces tests vérifient que le nébuliseur est capable de générer le flux requis)
        "Mesure Flux d'air (Réglage Bas)": {
            "injected": 2.5, # Exemple de référence de débit Faible (L/min)
            "tolerance": 1.0, 
            "unit": "L/min", 
            "type": "absolute", 
            "consigne": "Valeur mesurée doit être proche du réglage"
        },
        "Mesure Flux d'air (Réglage Moyen)": {
            "injected": 5.0, # Exemple de référence de débit Moyen (L/min)
            "tolerance": 1.5, 
            "unit": "L/min", 
            "type": "absolute", 
            "consigne": "Valeur mesurée doit être proche du réglage"
        },
        "Mesure Flux d'air (Réglage Haut)": {
            "injected": 8.0, # Exemple de référence de débit Élevé (L/min)
            "tolerance": 2.0, 
            "unit": "L/min", 
            "type": "absolute", 
            "consigne": "Valeur mesurée doit être proche du réglage"
        },
        
        # Temps de nébulisation (Test d'efficacité - Ex: nébulisation de 5mL. Tolérance: ± 3 min)
        "Temps de nébulisation pour 5mL": {
            "injected": 10.0, 
            "tolerance": 3.0, 
            "unit": "min", 
            "type": "absolute", 
            "consigne": "Moins de 15 minutes pour 5 mL"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Vérifier le fonctionnement de la nébulisation et du réglage d'intensité": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Passage automatique en mode batterie et fonctionnement sur batterie": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement des commandes (boutons, écran, voyants)": "Fonctionnel (Check)",
        "Fonctionnement de l'alarme de fluide bas": "Fonctionnel (Check)",
        "Fonctionnement de l'alarme de température excessive": "Fonctionnel (Check)",
        "Fonctionnement de l'alarme de coupure d'alimentation": "Fonctionnel (Check)",
        "Contrôle général des alarmes (visuelles et sonores)": "Alarmes fonctionnelles (Ok/Echoué/NA)"
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence et état de tous les accessoires (masques, tubulures, chambre d’humidification)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "Contrôle mécanique (système de fixation, roulettes et freins, si mobile)",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - Le patient est en contact avec les accessoires de nébulisation)
    "SECURITY_CHECKS": {
        # Résistance de terre (pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Accessoires de nébulisation/humidification = Partie appliquée Type BF)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Pousse-Seringue": {
        "NAME": "Pousse-seringue (Pompe à seringue)",
        "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
        "FABRICANT_LIST": [
            "Fresenius Kabi",
            "B. Braun",
            "Medtronic (Covidien)",
            "Smiths Medical",
            "Terumo",
            "Autres"
        ],
    
    # 1. Tests de Performance (Mesures critiques : Débit et Pression d'Occlusion)
    "PERFORMANCE_CHECKS": {
        
        # Débit Bas (Tolérance de +/- 2.5 mL/H basée sur 47.5-52.5)
        "Précision Débit (50.0 mL/H)": {
            "injected": 50.0, 
            "tolerance": 2.5, 
            "unit": "mL/H", 
            "type": "absolute", 
            "consigne": "50.0 mL/H ± 2.5 mL/H"
        },
        # Débit Haut (Tolérance de +/- 5.0 mL/H basée sur 94.9-104.9)
        "Précision Débit (99.9 mL/H)": {
            "injected": 99.9, 
            "tolerance": 5.0, 
            "unit": "mL/H", 
            "type": "absolute", 
            "consigne": "99.9 mL/H ± 5.0 mL/H"
        },
        
        # Pression d'Occlusion (Contre-pression) : Le seuil est la consigne d'alarme. Tolérance usuelle de ± 20 mmHg.
        "Pression d'Occlusion (Alarme 1 - 50.0 mL/H)": {
            "injected": 375, 
            "tolerance": 20, 
            "unit": "mmHg", 
            "type": "absolute", 
            "consigne": "375 mmHg ± 20 mmHg"
        },
        "Pression d'Occlusion (Alarme 2 - 99.9 mL/H)": {
            "injected": 375, 
            "tolerance": 20, 
            "unit": "mmHg", 
            "type": "absolute", 
            "consigne": "375 mmHg ± 20 mmHg"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Fonctionnement des commandes (boutons, écran)": "Fonctionnel (Check)",
        "Fonctionnement de la commande patient (PCAM, si équipé)": "Fonctionnel (Bon/Mauvais/NT)",
        "Limitation de dose correcte (PCAM, si équipé)": "Fonctionnel (Oui/Non)",
        "Fonctionnement de l'alarme d'occlusion": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement de l'alarme de fin de perfusion": "Fonctionnel (Check)",
        "Fonctionnement des alarmes de seringue (piston et corps)": "Fonctionnel (Check)",
        "Fonctionnement sur batterie et alarme batterie faible": "Fonctionnel (Check)",
        "Contrôle général des alarmes (visuelles et sonores)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil",
        "Présence de tous les accessoires (câble secteur, fixation)",
        "Visibilité et lisibilité des inscriptions et des voyants",
        "Fonctionnalité de la fixation du pousse-seringue",
        "Fonctionnalité de blocage de la seringue et fermeture du capot",
        "État général de l'appareil (boîtier, absence de dommages)",
    ],
    
    # 4. Tests de Sécurité Électrique (Type CF - Contact critique avec le patient via la ligne de perfusion)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Pièce appliquée de Type CF)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 10 μA (0.010 mA)", "unit": "mA", "type": "leakage"}, 
    }
    },
    "Défibrillateur": {
        # Informations Générales
        "NAME": "Défibrillateur (Manuel et/ou DSA)",
        "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
        "FABRICANT_LIST": [
            "Zoll",
            "Philips",
            "Medtronic/Physio-Control",
            "Schiller",
            "Autres"
        ],
        
        # 1. Tests de Performance (Mesures d'énergie délivrée et Temps de charge)
        "PERFORMANCE_CHECKS": {
            
            # Énergie délivrée (Tolérances basées sur les normes CEI 60601-2-4 et les seuils fournis)
            
            # Pour les basses énergies (5J), nous utilisons la tolérance relative (15% = 0.75J)
            "Énergie délivrée (Consigne 5 J)": {
                "injected": 5, 
                "tolerance": 15, 
                "unit": "J", 
                "type": "tolerance", 
                "consigne": "5 J ± 15%"
            },
            # Pour 20J, nous utilisons la tolérance relative (15% = 3J) pour conserver la cohérence
            "Énergie délivrée (Consigne 20 J)": {
                "injected": 20, 
                "tolerance": 15, 
                "unit": "J", 
                "type": "tolerance", 
                "consigne": "20 J ± 15%"
            },
            # Pour les moyennes et hautes énergies, nous utilisons la tolérance absolue standard (±4 J)
            "Énergie délivrée (Consigne 50 J)": {
                "injected": 50, 
                "tolerance": 4, 
                "unit": "J", 
                "type": "absolute", 
                "consigne": "50 J ± 4 J"
            },
            "Énergie délivrée (Consigne 200 J)": {
                "injected": 200, 
                "tolerance": 4, 
                "unit": "J", 
                "type": "absolute", 
                "consigne": "200 J ± 4 J"
            },
            "Énergie délivrée (Consigne 360 J)": {
                "injected": 360, 
                "tolerance": 4, 
                "unit": "J", 
                "type": "absolute", 
                "consigne": "360 J ± 4 J"
            },
            
            # Temps de charge (max 15s)
            "Temps de montée en charge max 360J (sur secteur)": {
                "injected": 15, 
                "tolerance": 0, 
                "unit": "s", 
                "type": "less_than", 
                "consigne": "Temps de charge maximal à 360J < 15 secondes"
            },
            "Temps de montée en charge max 4e choc (sur batterie)": {
                "injected": 15, 
                "tolerance": 0, 
                "unit": "s", 
                "type": "less_than", 
                "consigne": "Temps de charge maximal au 4ème choc sur batterie < 15 secondes"
            },
        },
        
        # 2. Tests Spécifiques (Contrôle de Fonctionnement, Logique et Sécurité)
        "SPECIFIC_CHECKS": {
            "Autotest de l'appareil": "Bon fonctionnement (Bon/Mauvais/NA)",
            "Passage automatique en mode batterie et fonctionnement sur batterie": "Fonctionnel (Bon/Mauvais/NA)",
            "Vérification de la capacité de la batterie à délivrer une série de 4 chocs": "Série de chocs réussie (Bon/Mauvais/NA)",
            "Fonctionnement de fin de charge (= choc prêt)": "Indication visuelle/sonore correcte (Bon/Mauvais/NA)",
            "Impossibilité de décharge sans les commandes appropriées": "Sécurité assurée (Bon/Mauvais/NA)",
            "Fonctionnement de la décharge de sécurité (décharge interne)": "Décharge interne fonctionnelle (Bon/Mauvais/NA)",
            "Vérification de reconnaissance du rythme (Mode Semi-automatique/DSA)": "Reconnaissance correcte (Sinusal -> pas de choc; FV -> choc) (Bon/Mauvais/NA)",
            "Fonctionnement des alarmes (Coupure alim, Délai charge trop long, Décharge auto)": "Toutes les alarmes sont fonctionnelles (Bon/Mauvais/NA)",
        },
        
        # 3. Contrôles Visuels (État Général, ECG et Enregistrement)
        "VISUAL_CHECKS": [
            "État de propreté de l'appareil et du ventilateur d’alimentation",
            "Présence de tous les accessoires (câble secteur, câbles ECG/capteurs patient, batterie, palettes/électrodes)",
            "État des palettes et câbles de défibrillation (intégrité, absence de défaut)",
            "Vérifier la lisibilité des affichages, des voyants et du clavier",
            "Vérification du bon défilement du papier (imprimante)",
            "Vérification de l'Impression et/ou enregistrement des paramètres patient",
            "Vérification que le tracé des courbes (ECG) est correct et cohérent",
            "État général de l'appareil",
        ],
        
        # 4. Tests de Sécurité Électrique (Basse Fréquence)
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
            "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
            # Le défibrillateur est un appareil de Type CF (Contact Cardiaque), le courant de fuite est donc critique.
            "Courant de fuite patient (Type CF)": {"limit": "< 10 μA (0.010 mA)", "unit": "mA", "type": "leakage"}, 
        }
    },
    "Bistouri Electrique " : {
    
        # Informations Générales
        "NAME": "Bistouri Électrique (Électrochirurgie)",
        "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
        "FABRICANT_LIST": [
            "Valleylab",
            "Erbe",
            "ConMed",
            "Autres"
        ],
        
        # 1. Tests de Performance (Mesures critiques)
        "PERFORMANCE_CHECKS": {
            
            # Courants de fuite Haute Fréquence (HF)
            # Basé sur Ref 029: Sortie Active < 150 mA sous 200 Ohm
            "Courant de fuite HF (Sortie Monopolaire Active)": {
                "injected": 150, 
                "tolerance": 0,
                "unit": "mA", 
                "type": "less_than",
                "consigne": "Doit être < 150 mA"
            },
            # Basé sur Ref 030: Sortie Dispersive < 150 mA sous 200 Ohm
            "Courant de fuite HF (Sortie Plaque Dispersive)": {
                "injected": 150, 
                "tolerance": 0,
                "unit": "mA", 
                "type": "less_than",
                "consigne": "Doit être < 150 mA"
            },
            # Basé sur Ref 031: Sortie Bipolaire < 60 mA sous 100 Ohm
            "Courant de fuite HF (Sortie Bipolaire)": {
                "injected": 60, 
                "tolerance": 0,
                "unit": "mA", 
                "type": "less_than",
                "consigne": "Doit être < 60 mA"
            },
    
            # Contrôle des Puissances (Exemple de points de mesure - à adapter selon modèle)
            # Utilisation d'une tolérance standard de +/- 20% pour les puissances de sortie.
            "Puissance de sortie (Mode Coupe - 50W)": {
                "injected": 50, 
                "tolerance": 20, 
                "unit": "W", 
                "type": "tolerance", 
                "consigne": "50W ± 20%"
            },
            "Puissance de sortie (Mode Coagulation - 50W)": {
                "injected": 50, 
                "tolerance": 20, 
                "unit": "W", 
                "type": "tolerance", 
                "consigne": "50W ± 20%"
            },
            "Puissance de sortie (Mode Bipolaire - 20W)": {
                "injected": 20, 
                "tolerance": 20, 
                "unit": "W", 
                "type": "tolerance", 
                "consigne": "20W ± 20%"
            },
        },
        
        # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité Interne)
        "SPECIFIC_CHECKS": {
            
            "Essais d'activation (par pédale et par commande manuelle)": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Contrôle de la sécurité plaque (système de surveillance de contact REM/CQM)": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Contrôle de départ automatique (si fonction présente)": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Fonctionnement sur batterie (si modèle équipé)": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Contrôle des temps d'émission": "Temps d'émission correct (Ok/Echoué/NA/NT)",
            "Contrôle des alarmes": "Toutes les alarmes (surchauffe, erreur, plaque) sont fonctionnelles (Ok/Echoué/NA/NT)",
        },
        
        # 3. Contrôles Visuels (État Général)
        "VISUAL_CHECKS": [
            "État de propreté de l'appareil et du ventilateur d’alimentation",
            "Présence de tous les accessoires (câble secteur, pièces à main, pédales)",
            "Vérification de la lisibilité et la visibilité des inscriptions",
            "Vérifier la fixation du bistouri et le fonctionnement des roulettes/freins (si chariot)",
            "État général de l'appareil",
        ],
        
        # 4. Tests de Sécurité Électrique (Basse Fréquence)
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
            "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
            "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"},
        }
    },
        "Stérilisateur - Autoclave" : {
        # Informations Générales
           "NAME": "Autoclave / Stérilisateur à vapeur",
            "TSE_REQUIRED": False, # Test de Sécurité Électrique requis
            "FABRICANT_LIST": [
                "Melag",
                "Tuttnauer",
                "Getinge",
                "Autres"
            ],
    
        # 1. Tests de Performance (Mesures critiques)
        "PERFORMANCE_CHECKS": {
            # Basé sur Ref 038: Vérifier la remontée de pression sur 10min
            "Test de Remontée de Pression (Test du Vide)": {
                "injected": 300, 
                "tolerance": 0,
                "unit": "mBar", 
                "type": "less_than", # La valeur mesurée doit être INFÉRIEURE à la limite
                "consigne": "Remontée de pression < 300 mBar sur 10 minutes"
            },
            
            # Le contrôle des cycles est qualitatif mais essentiel.
            # On peut ajouter le contrôle de la température/pression pour un cycle type 134°C.
            "Température de Stérilisation (134°C)": {
                "injected": 134, 
                "tolerance": 1, 
                "unit": "°C", 
                "type": "range", 
                "consigne": "134 °C ± 1 °C au plateau de stérilisation"
            },
        },
        
        # 2. Tests Spécifiques (Contrôle de Fonctionnement des Cycles et Sécurité)
        "SPECIFIC_CHECKS": {
            
            "Contrôle mécanique de l'appareil (Porte, vannes, casiers)": "Ouverture, fermeture et déplacement sont fonctionnels (Ok/Echoué)",
            "Réalisation de l'autotest": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Résultat du test du vide (Cycle)": "Cycle réussi (Ok/Echoué/NA/NT)",
            "Résultat du test de pénétration vapeur (Test Hélix/B&D)": "Test réussi (virage homogène et complet) (Ok/Echoué/NA/NT)",
            "Résultat du test d’un cycle de stérilisation normal (Charge sèche)": "Cycle réussi et charge sèche (Ok/Echoué/NA/NT)",
            "Résultat du test de sécurité (Coupure d’urgence)": "Arrêt sécurisé et affichage message d'erreur (Ok/Echoué/NA/NT)",
            "Fonctionnement de l'enregistrement et Impression des cycles": "Trace/rapport imprimé ou enregistré correctement (Ok/Echoué/NA/NT)",
            "Contrôle des alarmes": "Toutes les alarmes (Porte, Fin de cycle, Anomalies, Niveau d’eau) sont fonctionnelles (Ok/Echoué/NA/NT)",
        },
        
        # 3. Contrôles Visuels (État Général)
        "VISUAL_CHECKS": [
            "État de propreté de l'appareil (y compris cuve et filtres)",
            "Vérifier la présence de tous les accessoires (câble secteur, tuyaux, bonbonne d’eau, paniers…)",
            "État extérieur du dispositif et fonctionnement des mécanismes (fixation, absence de fuites, joints)",
            "État général de l'appareil",
        ],
        
        # 4. Tests de Sécurité Électrique (Basé sur vos standards)
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
            "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
            "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"},
        }
    },
    "Appareil de Photothérapie" : {
        # Informations Générales
        "NAME": "Appareil de Photothérapie",
        "TSE_REQUIRED": True, 
        "FABRICANT_LIST": [
            "Natus", "GE Healthcare", "Philips", "Draeger", "Autres"
        ],
        
        # 1. Tests de Performance (Basés sur les seuils de votre fiche)
        "PERFORMANCE_CHECKS": {
            # Ref 025: Puissance UV (Éclairement) - Seuil Bas : 16 Klux
            # La valeur mesurée doit être SUPÉRIEURE à la consigne.
            "Puissance UV (Éclairement à 200 mm)": {
                "injected": 16, 
                "tolerance": 0,
                "unit": "kLux", 
                "type": "more_than", 
                "consigne": "> 16 kLux (selon fiche Humatem)"
            },
            
            # Ref 024: Température - Seuil Haut : 30 °C
            # La valeur mesurée doit être INFÉRIEURE à la consigne.
            "Température au plan de couchage (à 10 cm)": {
                "injected": 30, 
                "tolerance": 0, 
                "unit": "°C", 
                "type": "less_than", 
                "consigne": "< 30 °C"
            },
            
            # Ajout du test standard d'irradiance comme point de contrôle supplémentaire.
            "Irradiance au centre (μW/cm²/nm)": {
                "injected": 30, "tolerance": 5, "unit": "µW/cm²/nm", "type": "range",
                "consigne": "30 ± 5 µW/cm²/nm (selon fabricant)"
            },
        },
        
        # 2. Tests Spécifiques (Contrôle de Fonctionnement)
        "SPECIFIC_CHECKS": {
            

            # Ref 019: Autotest
            "Vérification de l'autotest au démarrage ": "Bon fonctionnement (Résultat Bon/Moyen/Mauvais)",
            # Ref 020: Batterie
            "Fonctionnement sur batterie et autonomie ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 023: Compteur
            "Vérification du compteur des tubes UV ou des LED ": "Compteur OK (à noter ou < 2000h)",
            # Ref 026: Minuterie
            "Vérification du fonctionnement de la minuterie (5 min) ": "Cohérence avec le chronomètre (Ok/Echoué)",
            # Ref 027: Éclairage secondaire
            "Fonctionnalité de l'éclairage secondaire (si disponible) ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 030: Alarmes (Couverture des 028, 029)
            "Vérification des alarmes (Coupure alim., Fin de minuterie) ": "Les alarmes sonores et visuelles se déclenchent (Ok/Echoué)",
        },
        
        # 3. Contrôles Visuels (État Général - Ref 013 à 017)
        "VISUAL_CHECKS": [
            "État de propreté de l'appareil ",
            "Présence de tous les accessoires (câble secteur, roulettes, etc.) ",
            "Fonctionnalité des roulements et freins des roulettes ",
            "Fonctionnalité de la hauteur variable (si applicable) ",
            "État général de l'appareil (choc, fissure, afficheurs) )",
        ],
        
        # 4. Tests de Sécurité Électrique (Basé sur vos standards)
        # Réf 031: Test de sécurité électrique (Ok/Très Basse Tension de Sécurité/Echoué/NA)
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
            "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
            "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"},
        }
    },
    
    "Aspirateur à Mucosités (Succion)" : {
        # Informations Générales
        "NAME": "Aspirateur à Mucosités (Succion)",
        "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
        "FABRICANT_LIST": [
            "Devilbiss",
            "Medela",
            "Laerdal",
            "Autres"
        ],
        
        # 1. Tests de Performance (Contrôle des seuils de vide/dépression)
        "PERFORMANCE_CHECKS": {
            # Ref 029: Vérification de la dépression maximale
            # Consigne : > -0,15 BAR ou -100 mmHg. Nous allons utiliser l'unité mmHg.
            # Le vide mesuré doit être INFÉRIEUR (plus négatif) à -100 mmHg.
            
            # NOTE IMPORTANTE : Les valeurs de vide sont négatives. 
            # Pour une mesure facile, on peut tester la valeur ABSOLUE de la dépression.
            # Si on mesure la dépression absolue, la tolérance est : Dépression > 450 mmHg (0,6 BAR)
            
            "Dépression Maximale (Vide)": {
                "injected": 450, 
                "tolerance": 50, 
                "unit": "mmHg", 
                "type": "more_than", 
                "consigne": "> 450 mmHg (selon fabricant, ou 0.6 BAR)"
            },
            
            # Ref 025: Autonomie de la batterie
            "Autonomie Minimale de la Batterie": {
                "injected": 20, 
                "tolerance": 0, 
                "unit": "min", 
                "type": "more_than", 
                "consigne": "Autonomie > 20 minutes"
            },
        },
        
        # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité Interne)
        "SPECIFIC_CHECKS": {
            
            # Ref 020: Autotest
            "Autotest de l'appareil ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 024: Fonctionnement mécanique (Moteur, Pompe)
            "Fonctionnement mécanique de l'appareil ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 025: Mode batterie
            "Passage automatique en mode batterie ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 031 & 032: Aspiration/Dépression
            "Fonctionnement du système d'aspiration et de dépression ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 033: Sécurité Trop-Plein
            "Fonctionnement du système de trop-plein (Flotteur) ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 034 & 035: Double bocal (si applicable)
            "Fonctionnement du système à double bocal et de la transition ": "Fonctionnel (Ok/Echoué/NA/NT)",
            # Ref 036: Alarmes
            "Fonctionnement des alarmes (Bocal plein, tuyaux bouché, coupure alim.) ": "Fonctionnel (Ok/Echoué/NA/NT)",
        },
        
        # 3. Contrôles Visuels (État Général - Ref 013 à 017)
        "VISUAL_CHECKS": [
            "État de propreté de l'appareil et des accessoires (bocal, couvercle, filtre) ",
            "Présence de tous les accessoires (filtre, tubulures, coude, bocal) ",
            "État extérieur du dispositif et fonctionnement des mécanismes (joints, fixations, support) ",
            "État général du dispositif ",
        ],
        
        # 4. Tests de Sécurité Électrique (Basé sur vos standards - Ref 037)
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
            "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
            # L'aspirateur n'est pas un appareil de contact direct prolongé (type B/BF/CF) mais le test reste standard.
            "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"},
        }
},
    
    "Petit Matériel" : {
    # Informations Générales
    "NAME": "Petit Matériel (Générique, sans fiche dédiée)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique si l'appareil est électrique
    "FABRICANT_LIST": [
        "Divers",
        "A définir par le technicien",
        "NA"
    ],
    
    # 1. Tests de Performance (Contrôle qualitatif de la fonction principale)
    "PERFORMANCE_CHECKS": {
        # Ref 019: Le test est qualitatif ("Bon/Moyen/Mauvais")
        "Contrôle de performance du matériel (Fonction principale)": {
            "injected": 1, 
            "tolerance": 0, 
            "unit": "Qualitatif", 
            "type": "qualitatif", 
            "consigne": "La fonction principale est assurée (Ex: Mesure cohérente, chauffe, lumière ON)"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement)
    "SPECIFIC_CHECKS": {
        
        # Ref 017: Autotest
        "Autotest de l'appareil (si applicable) ": "Fonctionnel (Check)",
        # Ref 018: Fonctionnement général
        "Contrôle de fonctionnement de l'appareil ": "Bon fonctionnement général (Bon/Moyen/Mauvais)",
    },
    
    # 3. Contrôles Visuels (État Général - Ref 012 à 015)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil (Ref 012)",
        "Présence de tous les accessoires (câble secteur, tubulures, etc.) ",
        "Visibilité et lisibilité des inscriptions (marquage CE, fabricant, etc.) ",
        "État général de l'appareil (choc, fissure, intégrité) ",
    ],
    
    # 4. Tests de Sécurité Électrique (Basé sur vos standards - Ref 020)
    "SECURITY_CHECKS": {
        # Ces tests s'appliquent UNIQUEMENT si l'appareil est branché au secteur (si TSE_REQUIRED est True)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"},
    }
},
    "Bain Thermostaté / Étuve" : {
    # Informations Générales
    "NAME": "Bain Thermostaté / Étuve (Contrôle précis de température)",
    "TSE_REQUIRED": True, 
    "FABRICANT_LIST": [
        "Memmert",
        "Julabo",
        "Thermo Fisher Scientific",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures de température - Tolérance très serrée)
    "PERFORMANCE_CHECKS": {
        # Tolérance de +/- 0.1 °C par rapport à la consigne mesurée par le testeur
        "Relevé de température (Consigne 20 °C)": {"injected": 20, "tolerance": 0.1, "unit": "°C", "type": "range", "consigne": "20.0 ± 0.1 °C"},
        "Relevé de température (Consigne 40 °C)": {"injected": 40, "tolerance": 0.1, "unit": "°C", "type": "range", "consigne": "40.0 ± 0.1 °C"},
        "Relevé de température (Consigne 60 °C)": {"injected": 60, "tolerance": 0.1, "unit": "°C", "type": "range", "consigne": "60.0 ± 0.1 °C"},
        "Relevé de température (Consigne 80 °C)": {"injected": 80, "tolerance": 0.1, "unit": "°C", "type": "range", "consigne": "80.0 ± 0.1 °C"},
        "Relevé de température (Consigne 100 °C)": {"injected": 100, "tolerance": 0.1, "unit": "°C", "type": "range", "consigne": "100.0 ± 0.1 °C"},
        "Relevé de température (Consigne 120 °C)": {"injected": 120, "tolerance": 0.1, "unit": "°C", "type": "range", "consigne": "120.0 ± 0.1 °C"},
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité)
    "SPECIFIC_CHECKS": {
        
        "Autotest de l'appareil au démarrage": "Bon fonctionnement (Bon/Moyen/Mauvais)",
        "Vérification des boutons de commande, de l'écran tactile et de la lisibilité des voyants": "Fonctionnel (Bon/Moyen/Mauvais)",
        "Contrôle de la concordance des températures (consigne, affichée, mesurée)": "Concordance vérifiée (Ok/Echoué)",
        "Vérification des alarmes (sonde déconnectée, température excessive)": "Les alarmes sonores et visuelles sont fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil",
        "Contrôler la présence de tous les accessoires (câble secteur, chambres, capots, etc.)",
        "Vérification du fonctionnement mécanique (ouverture/fermeture de la porte ou du capot)",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Basé sur vos standards)
    "SECURITY_CHECKS": {
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"},
    }
},
    "Échographe" : {
    # Informations Générales
        "NAME": "Échographe (Système d'Imagerie Ultrasonore)",
        "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
        "FABRICANT_LIST": [
            "GE Healthcare",
            "Philips Healthcare",
            "Siemens Healthineers",
            "Mindray",
            "Autres"
        ],
        
        # 1. Tests de Performance (Pas de mesure quantitative standard sans fantôme)
        "PERFORMANCE_CHECKS": {}, # Les contrôles de performance sont qualitatifs (voir SPECIFIC_CHECKS)
        
        # 2. Tests Spécifiques (Contrôle de Fonctionnement, Image et Sécurité Interne)
        "SPECIFIC_CHECKS": {
            
            # Contrôles de Fonctionnement
            "Autotest de l'appareil au démarrage": "Bon fonctionnement (Bon/Moyen/Echoué/NT)",
            "Controle de la trackball / souris et clavier": "Fonctionnel (Bon/Moyen/Echoué/NT)",
            "Controle mécanique des sondes et des connecteurs (absence de fissures/défauts sur les câbles)": "Intégrité physique et électrique (Bon/Moyen/Echoué/NT)",
            "Fonctionnement du TGC (Time Gain Compensation)": "Contrôle de l'image dû à la TGC (Ok/Echoué/NT)",
            "Fonctionnement des modes BM et Doppler (Couleur, Pulsé, etc.)": "Fonctionnel (Ok/Echoué/NT)",
            "Fonctionnement du reprographe (Imprimante)": "Fonctionnel (Bon/Moyen/Echoué/NT)",
            "Controle des alarmes (système, sondes)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
            
            # Qualité d'image (Nécessite l'utilisation d'un fantôme de test ou d'une procédure de référence)
            "Contrôle du moniteur (échelle de gris, géométrie, contraste, luminosité, rétroéclairage)": "Affichage optimal (Bon/Moyen/Echoué/NT)",
            "Controle de la Zone Morte (proche de la sonde)": "Qualité de l'image immédiate (Bon/Moyen/Echoué/NT)",
            "Controle de la Profondeur d'exploration": "Capacité à atteindre la profondeur maximale (Bon/Moyen/Echoué/NT)",
            "Controle de la Résolution Axiale et Latérale": "Qualité de détail et de séparation des échos (Bon/Moyen/Echoué/NT)",
            "Controle de la Résolution en Contraste": "Capacité à différencier les tissus (Bon/Moyen/Echoué/NT)",
        },
        
        # 3. Contrôles Visuels (État Général et Accessoires)
        "VISUAL_CHECKS": [
            "État de propreté de l'appareil et du ventilateur d’alimentation",
            "Présence de tous les accessoires (sondes, cordon d’alimentation, gel...)",
            "Vérification de la lisibilité et visibilité des voyants et affichages",
            "Fonctionnalité des roulettes et des freins (si chariot)",
            "Fonctionnalité du bras de support des sondes/moniteur",
            "État général de l'appareil",
        ],
        
        # 4. Tests de Sécurité Électrique (Basse Fréquence)
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
            "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
            # Les sondes sont des parties appliquées (Type BF)
            "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
        }
    },
    "Garrot Pneumatique" : {
    # Informations Générales
    "NAME": "Garrot Pneumatique (Tourniquet)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Zimmer",
        "Stryker",
        "VBM Medizintechnik",
        "A.T.S. (Anatomical Tourniquet System)",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Pression et Temps)
    "PERFORMANCE_CHECKS": {
        
        # Pression (Tolérance standard de +/- 10 mmHg pour la vérification)
        "Précision de la pression (Consigne 100 mmHg)": {
            "injected": 100, 
            "tolerance": 10, 
            "unit": "mmHg", 
            "type": "absolute", 
            "consigne": "100 mmHg ± 10 mmHg"
        },
        "Précision de la pression (Consigne 300 mmHg)": {
            "injected": 300, 
            "tolerance": 10, 
            "unit": "mmHg", 
            "type": "absolute", 
            "consigne": "300 mmHg ± 10 mmHg"
        },
        "Précision de la pression (Consigne 500 mmHg)": {
            "injected": 500, 
            "tolerance": 10, 
            "unit": "mmHg", 
            "type": "absolute", 
            "consigne": "500 mmHg ± 10 mmHg"
        },
        
        # Minuterie (Tolérance de 5% pour la vérification du temps)
        "Contrôle de la minuterie (10 min)": {
            "injected": 10, 
            "tolerance": 5, 
            "unit": "min", 
            "type": "tolerance", 
            "consigne": "10 minutes ± 5%"
        },
        "Contrôle de la minuterie (60 min)": {
            "injected": 60, 
            "tolerance": 5, 
            "unit": "min", 
            "type": "tolerance", 
            "consigne": "60 minutes ± 5%"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Fonctionnement du tableau de commande et de l'affichage": "Fonctionnel (Check)",
        "Passage automatique en mode batterie et fonctionnalité (incluant détection faible)": "Fonctionnel (Check)",
        "Contrôle de la pression (stabilité courte et longue durée)": "Maintien de la pression (Ok/Echoué/NA/NT)",
        "Vérification des alarmes (fuite, occlusion, batterie faible)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur (si présent)",
        "Présence de tous les accessoires (Brassards, tubes de raccordements)",
        "Aspect général du garrot et des brassards (pas de fissures ou dommages)",
        "Lisibilité des affichages et des étiquettes",
        "Sécurité mécanique de l'appareil (système de fixation, support)",
        "Contrôles visuels et mécaniques généraux",
    ],
    
    # 4. Tests de Sécurité Électrique (Basse Fréquence)
    "SECURITY_CHECKS": {
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Le brassard de garrot pneumatique n'est pas une partie appliquée au sens strict (pas de contact électrique direct avec le patient).
        "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Humidificateur-Réchauffeur" : {
    # Informations Générales
    "NAME": "Humidificateur-Réchauffeur",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Fisher & Paykel",
        "ResMed",
        "Maquet/Getinge",
        "Intersurgical",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Température)
    "PERFORMANCE_CHECKS": {
        # Basé sur le test 022/023 : concordance Température de consigne vs Température mesurée
        "Précision de la température (Écart maximal)": {
            "injected": 37, # Consigne typique pour circuit patient
            "tolerance": 1.5, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "Écart Temp. Consigne / Temp. Mesurée ≤ 1,5 °C"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement, Sondes et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement des commandes (boutons, écran, voyants, fonctionnement sonore)": "Fonctionnel (Ok/Echoué/NA/NT)",
        # Basé sur le test 025 : Précision des sondes par rapport entre elles et au testeur
        "Contrôle des sondes (Écart interne et Écart au testeur ≤ 1,5 °C)": "Précision des sondes (Ok/Echoué/NA/NT)",
        # Basé sur les tests 027 à 031
        "Contrôle des alarmes (Sonde déconnectée, Basse Temp., Haute Temp., Fil chauffant déconnecté)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        "Contrôle du fonctionnement mécanique global (incl. support chambre humidification)": "Fonctionnel (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence de tous les accessoires (câble secteur, fil chauffant, chambre d’humidification, embouts de tuyaux…)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Basse Fréquence)
    "SECURITY_CHECKS": {
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Partie appliquée de Type BF (via le circuit patient)
        "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Incubateur de Néonatalogie Fermé" : {
    # Informations Générales
        "NAME": "Incubateur de Néonatalogie Fermé",
        "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
        "FABRICANT_LIST": [
            "Dräger",
            "GE Healthcare",
            "Atom",
            "Natus",
            "Autres"
        ],
        
        # 1. Tests de Performance (Mesures critiques : Température, Humidité, O2, Bruit)
        "PERFORMANCE_CHECKS": {
            
            # Température de l'air (Tolérance standard: ± 1,0 °C)
            "Précision de la température de l'air (36 °C)": {
                "injected": 36.0, 
                "tolerance": 1.0, 
                "unit": "°C", 
                "type": "absolute", 
                "consigne": "36.0 °C ± 1.0 °C"
            },
            # Température cutanée (Tolérance standard: ± 0,5 °C à ± 1,0 °C)
            "Précision de la température cutanée (36.5 °C)": {
                "injected": 36.5, 
                "tolerance": 1.0, 
                "unit": "°C", 
                "type": "absolute", 
                "consigne": "36.5 °C ± 1.0 °C"
            },
            
            # Hygrométrie (Tolérance standard: ± 5% RH)
            "Précision de l'hygrométrie (60% RH)": {
                "injected": 60, 
                "tolerance": 5, 
                "unit": "% RH", 
                "type": "absolute", 
                "consigne": "60% RH ± 5% RH"
            },
            
            # Oxymétrie (Tolérance standard: ± 2% O2)
            "Précision de l'Oxymétrie (50% O2)": {
                "injected": 50, 
                "tolerance": 2, 
                "unit": "% O2", 
                "type": "absolute", 
                "consigne": "50% O2 ± 2% O2"
            },
    
            # Niveau Sonore (Doit être inférieur à 60 dBa)
            "Niveau sonore dans l'enceinte (< 60 dBa)": {
                "injected": 58, 
                "tolerance": 2, 
                "unit": "dBa", 
                "type": "absolute", 
                "consigne": "Mesure ≤ 60 dBa (hors alarmes)"
            },
        },
        
        # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité Interne)
        "SPECIFIC_CHECKS": {
            "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
            "Autotest de l'appareil au démarrage": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Fonctionnement des commandes, écran, voyants et alarme sonore": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Fonctionnement des ventilateurs (principal et supplémentaire)": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Vérification du fonctionnement sur batterie": "Fonctionnel (Ok/Echoué/NA/NT)",
            "Contrôle de la pression des fluides gazeux (O2, air, vide)": "Pression conforme (Ok/Echoué/NA/NT)",
            "Contrôle d'étanchéité de l'enceinte et des valves": "Absence de fuite (Ok/Echoué/NA/NT)",
            "Contrôle des alarmes (Température, Sondes, Alimentation, O2)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        },
        
        # 3. Contrôles Visuels (État Général et Mécanique)
        "VISUAL_CHECKS": [
            "État de propreté de l'appareil, du ventilateur d’alimentation et de l'enceinte",
            "Contrôle de la date du filtre à air et remplacement si nécessaire",
            "Présence de tous les accessoires (sondes, tuyaux, etc.)",
            "Vérification du code des couleurs et raccordement des tuyaux",
            "Vérification de la bonne fonctionnalité des éléments mécaniques (porte, hublots, freins, verrouillages châssis, inclinaisons)",
            "Lisibilité des inscriptions et des affichages",
            "État général de l'appareil",
        ],
        
        # 4. Tests de Sécurité Électrique (Type BF - Parties appliquées)
        "SECURITY_CHECKS": {
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
            "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
            # Les sondes cutanées sont des parties appliquées de Type BF
            "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
        }
    },
    "Incubateur de Néonatalogie Ouvert " : {
    # Informations Générales
    "NAME": "Incubateur de Néonatalogie Ouvert (Table de Réanimation)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Dräger",
        "GE Healthcare",
        "Atom",
        "Natus",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Température, Humidité, O2)
    "PERFORMANCE_CHECKS": {
        
        # Température cutanée (Tolérance standard: ± 1,0 °C)
        "Précision de la Température Cutanée (36.5 °C)": {
            "injected": 36.5, 
            "tolerance": 1.0, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "36.5 °C ± 1.0 °C"
        },
        # Température du Matelas Chauffant (Tolérance de ± 1,0 °C pour le mode chauffage)
        "Précision de la Température du Matelas (38.0 °C)": {
            "injected": 38.0, 
            "tolerance": 1.0, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "38.0 °C ± 1.0 °C"
        },
        
        # Hygrométrie (Si équipé, Tolérance standard: ± 5% RH)
        "Précision de l'hygrométrie (60% RH)": {
            "injected": 60, 
            "tolerance": 5, 
            "unit": "% RH", 
            "type": "absolute", 
            "consigne": "60% RH ± 5% RH"
        },
        
        # Oxymétrie (Si équipé, Tolérance standard: ± 2% O2)
        "Précision de l'Oxymétrie (50% O2)": {
            "injected": 50, 
            "tolerance": 2, 
            "unit": "% O2", 
            "type": "absolute", 
            "consigne": "50% O2 ± 2% O2"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité Interne)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement des commandes (boutons, écran, voyants) et fonctionnement sonore": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement du ventilateur supplémentaire (si présent)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Vérification du fonctionnement sur batterie/accumulateur d'alarme": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Contrôle de pression des fluides gazeux (O2, air, vide)": "Pression conforme (Ok/Echoué/NA/NT)",
        "Contrôle d'étanchéité des valves des fluides gazeux": "Absence de fuite (Ok/Echoué/NA/NT)",
        "Contrôle des alarmes (Coupure alim, Sondes déconnectées, Température excessive)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence de tous les accessoires (sondes, tuyaux, fil chauffant, etc.)",
        "Vérification de la date de l’accumulateur pour l’alarme de coupure secteur",
        "Vérification du respect du code des couleurs et raccordement des tuyaux",
        "Vérification de la bonne fonctionnalité des éléments mécaniques (barrières latérales, freins, verrouillages châssis, inclinaisons)",
        "Lisibilité des inscriptions et des affichages",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - Parties appliquées)
    "SECURITY_CHECKS": {
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Les sondes cutanées sont des parties appliquées de Type BF
        "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Moniteur d’Oxygène" : {
    # Informations Générales
    "NAME": "Moniteur d’Oxygène",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (si sur secteur)
    "FABRICANT_LIST": [
        "Dräger",
        "Maxtec",
        "Vyaire Medical",
        "Pac-Sci",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Concentration en O2)
    "PERFORMANCE_CHECKS": {
        
        # Test de précision à haute concentration (basé sur Seuil Bas 97 / Seuil Haut 103)
        "Précision de l'O2 (Consigne 100%)": {
            "injected": 100, 
            "tolerance": 3, 
            "unit": "% O2", 
            "type": "absolute", 
            "consigne": "100% O2 ± 3%"
        },
        
        # Test de précision à concentration ambiante (Air Ambiant)
        "Précision de l'O2 (Air Ambiant 21%)": {
            "injected": 21, # Concentration standard de l'air ambiant
            "tolerance": 1, 
            "unit": "% O2", 
            "type": "absolute", 
            "consigne": "21% O2 ± 1%"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement des commandes (boutons, écran, voyants)": "Fonctionnel (Check)",
        "Fonctionnement sur pile/batterie (incluant le passage automatique si applicable)": "Autonomie fonctionnelle (Ok/Echoué/NA/NT)",
        "Contrôle de l’étalonnage du capteur (si en dehors de la tolérance)": "Calibré (Check)",
        "Fonctionnement des alarmes (Seuil Haut/Bas, Batterie faible, Capteur défectueux)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil, du ventilateur d’alimentation et du capteur",
        "Présence de tous les accessoires (câble secteur, capteur, raccords, etc.)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "Contrôle du système de fixation (si présent)",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Type B ou sans partie appliquée)
    "SECURITY_CHECKS": {
        # La résistance de terre est pertinente si l'appareil est de Classe I (sur secteur)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (pertinent pour Classe I et II)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Le capteur O2 peut être considéré comme partie appliquée (Type BF) s'il est utilisé en circuit fermé avec le patient.
        "Courant de fuite patient": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Moniteur ECG" : {
    # Informations Générales
    "NAME": "Moniteur ECG",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Philips",
        "GE Healthcare",
        "Mindray",
        "Schiller",
        "Masimo",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Fréquence Cardiaque, Respiration, Température)
    "PERFORMANCE_CHECKS": {
        
        # Précision du Pouls (ECG) - Bas (Tolérance: ± 5 bpm)
        "Précision Pouls ECG (Basse - 30 bpm)": {
            "injected": 30, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "30 bpm ± 5 bpm"
        },
        # Précision du Pouls (ECG) - Moyen (Tolérance: ± 5 bpm)
        "Précision Pouls ECG (Moyenne - 120 bpm)": {
            "injected": 120, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "120 bpm ± 5 bpm"
        },
        # Précision du Pouls (ECG) - Haut (Tolérance: ± 5 bpm)
        "Précision Pouls ECG (Haute - 240 bpm)": {
            "injected": 240, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "240 bpm ± 5 bpm"
        },

        # Respiration (Tolérance standard: ± 5 resp/min)
        "Précision Fréquence Respiratoire (30 resp/min)": {
            "injected": 30, 
            "tolerance": 5, 
            "unit": "resp/min", 
            "type": "absolute", 
            "consigne": "30 resp/min ± 5 resp/min"
        },
        
        # Température (Tolérance standard: ± 0.2 °C)
        "Précision de la Température (37.0 °C)": {
            "injected": 37.0, 
            "tolerance": 0.2, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "37.0 °C ± 0.2 °C"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Contrôle de la qualité du signal ECG (ligne de base, forme des courbes)": "Signal clair (Check)",
        "Passage automatique et fonctionnement sur batterie": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement des commandes (boutons, écran tactile, voyants)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Vérification de l'enregistrement et de l'impression des paramètres": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Contrôle des alarmes ECG (FC Basse/Haute, Arythmie, Câble)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        "Contrôle des alarmes Respiration (FR Basse/Haute, Apnée)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        "Contrôle des alarmes Température (Basse/Haute, Sonde)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence et état de tous les accessoires (câbles ECG, sondes de température, etc.)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "Contrôle mécanique (fixation, fluidité des roulettes et freins)",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - Le module ECG nécessite souvent Type CF)
    "SECURITY_CHECKS": {
        # Résistance de terre (pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Type BF ou CF, utiliser BF comme minimum ici)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Moniteur Multiparamétrique" : {
    # Informations Générales
    "NAME": "Moniteur Multiparamétrique",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Philips",
        "GE Healthcare",
        "Mindray",
        "Dräger",
        "Masimo",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques pour chaque module)
    "PERFORMANCE_CHECKS": {
        
        # Température (Utilisation d'une valeur clinique de 37°C avec tolérance de +/- 0.2°C)
        "Précision de la Température (37.0 °C)": {
            "injected": 37.0, 
            "tolerance": 0.2, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "37.0 °C ± 0.2 °C"
        },
        
        # SpO2 (Tolérance de ± 3% basée sur les seuils fournis)
        "Précision SpO2 (Basse - 80%)": {
            "injected": 80, 
            "tolerance": 3, 
            "unit": "% SpO2", 
            "type": "absolute", 
            "consigne": "80% ± 3%"
        },
        "Précision SpO2 (Intermédiaire - 93%)": {
            "injected": 93, 
            "tolerance": 3, 
            "unit": "% SpO2", 
            "type": "absolute", 
            "consigne": "93% ± 3%"
        },
        "Précision SpO2 (Haute - 97%)": {
            "injected": 97, 
            "tolerance": 3, 
            "unit": "% SpO2", 
            "type": "absolute", 
            "consigne": "97% ± 3%"
        },
        
        # ECG / Fréquence Cardiaque (FC) (Tolérance de ± 5 bpm)
        "Précision FC (Basse - 30 bpm)": {
            "injected": 30, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "30 bpm ± 5 bpm"
        },
        "Précision FC (Intermédiaire - 120 bpm)": {
            "injected": 120, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "120 bpm ± 5 bpm"
        },
        "Précision FC (Haute - 240 bpm)": {
            "injected": 240, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "240 bpm ± 5 bpm"
        },
        
        # Respiration (Tolérance de ± 3 resp/min)
        "Précision Respiration (Basse - 15 resp/min)": {
            "injected": 15, 
            "tolerance": 3, 
            "unit": "resp/min", 
            "type": "absolute", 
            "consigne": "15 resp/min ± 3 resp/min"
        },
        "Précision Respiration (Haute - 30 resp/min)": {
            "injected": 30, 
            "tolerance": 3, 
            "unit": "resp/min", 
            "type": "absolute", 
            "consigne": "30 resp/min ± 3 resp/min"
        },
        
        # PNI / NIBP (Contrôle à différentes pressions. Tolérance de ± 8 mmHg pour Sys/Dia/MAP)
        "Précision PNI (Basse - 80/40 mmHg)": {
            "injected": 80, 
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique ± 8 mmHg"
        },
        "Précision PNI (Moyenne - 120/80 mmHg)": {
            "injected": 120, 
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique ± 8 mmHg"
        },
        "Précision PNI (Haute - 180/140 mmHg)": {
            "injected": 180, 
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique ± 8 mmHg"
        },
        "Précision PNI (Néonat - 70/40 mmHg)": {
            "injected": 70, 
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique Néonat ± 8 mmHg"
        },
        
        # EtCO2 (Tolérance de ± 2 mmHg pour le CO2)
        "Précision EtCO2 (35 mmHg)": {
            "injected": 35, 
            "tolerance": 2, 
            "unit": "mmHg", 
            "type": "absolute", 
            "consigne": "35 mmHg (4.7% CO2) ± 2 mmHg"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Passage automatique en mode batterie et autonomie": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement des commandes (boutons, écran tactile, voyants)": "Fonctionnel (Check)",
        "Qualité du signal ECG (forme d'onde PQRS, PVC)": "Bonne allure des courbes (Bon/Mauvais)",
        "Cohérence du nombre de cycles respiratoires (EtCO2)": "Cohérent (Check)",
        "Enregistrement et impression des paramètres patient (si équipé)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Contrôle global des alarmes (visuelles et sonores) pour tous les modules": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence et état de tous les accessoires (câbles ECG, capteur SpO2, brassards PNI, sondes Temp/EtCO2, etc.)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "Contrôle mécanique du moniteur (fixation sur le pied, roulements et freins des roulettes)",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF/CF selon les modules)
    "SECURITY_CHECKS": {
        # Résistance de terre (pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Les modules ECG et PNI sont des parties appliquées Type CF ou BF)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Moniteur PNI" : {
    # Informations Générales
    "NAME": "Moniteur PNI (Pression Artérielle Non Invasive)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Welch Allyn",
        "GE Healthcare",
        "Mindray",
        "Omron",
        "Bionet",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Pression Artérielle)
    "PERFORMANCE_CHECKS": {
        
        # PNI (Tolérance de ± 8 mmHg pour les pressions Systolique et Diastolique)
        "Précision PNI (Basse - 80/40 mmHg)": {
            "injected": 80, # Référence Systolique
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique ± 8 mmHg"
        },
        "Précision PNI (Moyenne - 120/80 mmHg)": {
            "injected": 120, 
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique ± 8 mmHg"
        },
        "Précision PNI (Haute - 180/140 mmHg)": {
            "injected": 180, 
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique ± 8 mmHg"
        },
        "Précision PNI (Néonat - 70/40 mmHg)": {
            "injected": 70, 
            "tolerance": 8, 
            "unit": "mmHg (Sys)", 
            "type": "absolute", 
            "consigne": "Systolique et Diastolique Néonat ± 8 mmHg"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Fonctionnement des commandes (boutons, écran tactile, voyants)": "Fonctionnel (Check)",
        "Contrôle de l'étanchéité du circuit et du gonflage maximal (sur-pression)": "Conforme (Check)",
        "Contrôle des alarmes PNI (Seuil Bas/Haut, brassard déconnecté, fuite)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        "Enregistrement et impression des paramètres patient (si équipé)": "Fonctionnel (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence et état de tous les accessoires (câble secteur, tubes PNI, brassards Adulte/Néonat)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "Contrôle mécanique (fixation sur le pied, roulements et freins des roulettes)",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - Le brassard PNI est une partie appliquée)
    "SECURITY_CHECKS": {
        # Résistance de terre (pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Brassard PNI = Partie appliquée Type BF)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Moniteur SpO2" : {
    # Informations Générales
    "NAME": "Moniteur SpO2 (Oxymétrie de Pouls)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Masimo",
        "Nellcor (Medtronic)",
        "Nonin Medical",
        "GE Healthcare",
        "Philips",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Saturation en O2)
    "PERFORMANCE_CHECKS": {
        
        # SpO2 Basse (cible 80%, tolérance ± 3% basée sur 77-83)
        "Précision SpO2 (Basse - 80%)": {
            "injected": 80, 
            "tolerance": 3, 
            "unit": "% SpO2", 
            "type": "absolute", 
            "consigne": "80% ± 3%"
        },
        # SpO2 Intermédiaire (cible 93%, tolérance ± 3% basée sur 90-96)
        "Précision SpO2 (Intermédiaire - 93%)": {
            "injected": 93, 
            "tolerance": 3, 
            "unit": "% SpO2", 
            "type": "absolute", 
            "consigne": "93% ± 3%"
        },
        # SpO2 Haute (cible 97%, tolérance ± 3% basée sur 94-100)
        "Précision SpO2 (Haute - 97%)": {
            "injected": 97, 
            "tolerance": 3, 
            "unit": "% SpO2", 
            "type": "absolute", 
            "consigne": "97% ± 3%"
        },
        
        # Fréquence Cardiaque (Pouls) – Test de vérification du simulateur de SpO2. Cible 120 bpm (standard)
        "Précision Pouls SpO2 (120 bpm)": {
            "injected": 120, 
            "tolerance": 5, 
            "unit": "bpm", 
            "type": "absolute", 
            "consigne": "120 bpm ± 5 bpm"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Fonctionnement des commandes (boutons, écran, voyants)": "Fonctionnel (Check)",
        "Fonctionnement sur pile/batterie (si équipé)": "Autonomie fonctionnelle (Check)",
        "Contrôle des alarmes SpO2 (Seuil Bas/Haut, sonde déconnectée)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        "Enregistrement et impression des paramètres patient (si équipé)": "Fonctionnel (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence et état de tous les accessoires (câble secteur, sondes SpO2 de différents types)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "Contrôle mécanique (fixation sur le pied, roulements et freins des roulettes, si mobile)",
        "État général de l'appareil",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - La sonde SpO2 est une partie appliquée)
    "SECURITY_CHECKS": {
        # Résistance de terre (pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Sonde SpO2 = Partie appliquée Type BF)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Pèse bébé" : {
    # Informations Générales
    "NAME": "Pèse bébé",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis
    "FABRICANT_LIST": [
        "Seca",
        "Charder",
        "Tanita",
        "Detecto",
        "Health o meter",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Précision de la pesée)
    "PERFORMANCE_CHECKS": {
        
        # Précision (Tolérance de ± 10 g basée sur les seuils fournis)
        "Précision (500 g)": {
            "injected": 500, 
            "tolerance": 10, 
            "unit": "g", 
            "type": "absolute", 
            "consigne": "500 g ± 10 g"
        },
        "Précision (1000 g)": {
            "injected": 1000, 
            "tolerance": 10, 
            "unit": "g", 
            "type": "absolute", 
            "consigne": "1000 g ± 10 g"
        },
        "Précision (1500 g)": {
            "injected": 1500, 
            "tolerance": 10, 
            "unit": "g", 
            "type": "absolute", 
            "consigne": "1500 g ± 10 g"
        },
        "Précision (2000 g)": {
            "injected": 2000, 
            "tolerance": 10, 
            "unit": "g", 
            "type": "absolute", 
            "consigne": "2000 g ± 10 g"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Contrôle de la tare à 0": "Fonctionnel (Check)",
        "Fonctionnement des commandes (boutons, écran)": "Fonctionnel (Check)",
        "Vérifier le fonctionnement sur batterie (> 30 min)": "Autonomie fonctionnelle (Bon/Mauvais/NA/NT)",
        "Vérifier la lisibilité des voyants (secteur/batterie)": "Fonctionnel (Check)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil",
        "Présence et état de tous les accessoires (plateau amovible, adaptateur secteur)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "État général de l'appareil (absence de fissures, plateau stable)",
    ],
    
    # 4. Tests de Sécurité Électrique (Type B - Le bébé est en contact avec la partie appliquée)
    "SECURITY_CHECKS": {
        # Résistance de terre (pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Plateau de pesée = Partie appliquée Type B)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Poupinel / Étuve" : {
    # Informations Générales
    "NAME": "Poupinel / Étuve (Stérilisateur à chaleur sèche)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I)
    "FABRICANT_LIST": [
        "Memmert",
        "Binder",
        "Steri-Max",
        "VWR",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Précision de la Température)
    "PERFORMANCE_CHECKS": {
        
        # Température (Tolérance de ± 5 °C basée sur les seuils fournis)
        "Précision Température (50 °C)": {
            "injected": 50, 
            "tolerance": 5, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "50 °C ± 5 °C"
        },
        "Précision Température (100 °C)": {
            "injected": 100, 
            "tolerance": 5, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "100 °C ± 5 °C"
        },
        "Précision Température (150 °C)": {
            "injected": 150, 
            "tolerance": 5, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "150 °C ± 5 °C"
        },
        "Précision Température (200 °C)": {
            "injected": 200, 
            "tolerance": 5, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "200 °C ± 5 °C"
        },
        "Précision Température (250 °C)": {
            "injected": 250, 
            "tolerance": 5, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "250 °C ± 5 °C"
        },
        "Précision Température (300 °C)": {
            "injected": 300, 
            "tolerance": 5, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "300 °C ± 5 °C"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil": "Fonctionnel (Bon/Moyen/Mauvais)",
        "Fonctionnement des commandes, écran tactile et lisibilité des voyants": "Fonctionnel (Bon/Moyen/Mauvais)",
        "Fonctionnement mécanique de la porte (ouverture/fermeture)": "Fonctionnel (Check)",
        "Vérification du fonctionnement sonore": "Fonctionnel (Check)",
        "Vérification des alarmes (sonde déconnectée/température excessive)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        "Contrôle de la température général (conformité des 6 relevés)": "Conforme (Ok/Echoué/NA/NT)"
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence et état de tous les accessoires (paniers, étagères, câble secteur)",
        "Visibilité et lisibilité des inscriptions et des étiquettes",
        "État général de l'appareil (isolation, joint de porte, intérieur de l'enceinte)",
    ],
    
    # 4. Tests de Sécurité Électrique (Appareil de Classe I)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient : Non Applicable (Pas de contact patient direct)
        "Courant de fuite patient (parties appliquées)": {"limit": "NA", "unit": "mA", "type": "leakage"}, 
    }
},
    "Ventilateur de réanimation / urgence" : {
    # Informations Générales
    "NAME": "Ventilateur de réanimation / urgence",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I)
    "FABRICANT_LIST": [
        "Dräger",
        "Hamilton Medical",
        "Maquet/Getinge",
        "Vyaire Medical",
        "Philips",
        "Mindray",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : FiO2, Volume, Pression, Fréquence)
    # Note : Les tolérances sont fixées à ±10% basées sur les seuils fournis.
    "PERFORMANCE_CHECKS": {
        
        # --- Contrôle en Volume (Volume Control) ---
        "VC - FiO2 (21 %)": {
            "injected": 21, 
            "tolerance": 10, 
            "unit": "%", 
            "type": "percentage", 
            "consigne": "21 % ± 10 %"
        },
        "VC - Volume Minute (5 L)": {
            "injected": 5, 
            "tolerance": 10, 
            "unit": "L", 
            "type": "percentage", 
            "consigne": "5 L ± 10 %"
        },
        "VC - Volume Courant (500 mL)": {
            "injected": 500, 
            "tolerance": 10, 
            "unit": "mL", 
            "type": "percentage", 
            "consigne": "500 mL ± 10 %"
        },
        "VC - Pression Maximale (40 cmH2O)": {
            "injected": 40, 
            "tolerance": 10, 
            "unit": "cmH2O", 
            "type": "percentage", 
            "consigne": "40 cmH2O ± 10 %"
        },
        "VC - Fréquence (10 c/min)": {
            "injected": 10, 
            "tolerance": 10, 
            "unit": "c/min", 
            "type": "percentage", 
            "consigne": "10 c/min ± 10 %"
        },
        "VC - Ratio I/E (1:2 soit 0.5)": {
            "injected": 0.5, 
            "tolerance": 10, 
            "unit": "", 
            "type": "percentage", 
            "consigne": "0.5 ± 10 %"
        },
        "VC - PEEP (4 cmH2O)": {
            "injected": 4, 
            "tolerance": 10, 
            "unit": "cmH2O", 
            "type": "percentage", 
            "consigne": "4 cmH2O ± 10 %"
        },
        
        # --- Contrôle en Pression (Pressure Control) ---
        "PC - FiO2 (21 %)": {
            "injected": 21, 
            "tolerance": 10, 
            "unit": "%", 
            "type": "percentage", 
            "consigne": "21 % ± 10 %"
        },
        "PC - Pression Inspiratoire (30 cmH2O)": { # Basé sur la consigne mesurée fournie (30 ± 3)
            "injected": 30, 
            "tolerance": 10, 
            "unit": "cmH2O", 
            "type": "percentage", 
            "consigne": "30 cmH2O ± 10 %"
        },
        "PC - Fréquence (10 c/min)": {
            "injected": 10, 
            "tolerance": 10, 
            "unit": "c/min", 
            "type": "percentage", 
            "consigne": "10 c/min ± 10 %"
        },
        "PC - Ratio I/E (1:2 soit 0.5)": {
            "injected": 0.5, 
            "tolerance": 10, 
            "unit": "", 
            "type": "percentage", 
            "consigne": "0.5 ± 10 %"
        },
        "PC - PEEP (4 cmH2O)": {
            "injected": 4, 
            "tolerance": 10, 
            "unit": "cmH2O", 
            "type": "percentage", 
            "consigne": "4 cmH2O ± 10 %"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Fonctionnement de l'Autotest au démarrage": "Fonctionnel (Check)",
        "Fonctionnement sur batterie (capacité et charge)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Vérification des tests de fuite internes à l'appareil": "Conforme (Ok/Echoué/NA/NT)",
        "Étalonnage de la cellule à O2 (21 %)": "Conforme (Ok/Echoué/NA/NT)",
        "Étalonnage de la cellule à O2 (100 % - si gaz disponible)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement du Trigger d'inspiration (-5 cmH2O)": "Fonctionnel (Ok/Echoué)",
        "Contrôle global des alarmes (gaz, pression, alim., fuite)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
        "Contrôle mécanique (roulettes, freins, bras, fixations)": "Conforme (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence de tous les accessoires (capteurs, tuyaux gaz/patient, filtre, humidification)",
        "Visibilité et lisibilité des inscriptions et des voyants",
        "État général de l'appareil (boîtier, écran, raccords)",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - contact patient par les circuits de respiration)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Partie appliquée de Type BF - circuit de respiration)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Centrifugeuse (Laboratoire/Clinique)" : {
    # Informations Générales
    "NAME": "Centrifugeuse (Laboratoire/Clinique)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I)
    "FABRICANT_LIST": [
        "Hettich",
        "Eppendorf",
        "Thermo Scientific",
        "Beckman Coulter",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Tachymétrie et Température)
    "PERFORMANCE_CHECKS": {
        
        # --- Tachymétrie (Vitesse de Rotation) : Tolérance de +/- 200 RPM (tours/minute) ---
        "Précision Tachymétrie (15000 RPM)": {
            "injected": 15000, 
            "tolerance": 200, 
            "unit": "RPM", 
            "type": "absolute", 
            "consigne": "15000 RPM ± 200 RPM"
        },
        "Précision Tachymétrie (10000 RPM)": {
            "injected": 10000, 
            "tolerance": 200, 
            "unit": "RPM", 
            "type": "absolute", 
            "consigne": "10000 RPM ± 200 RPM"
        },
        "Précision Tachymétrie (5000 RPM)": {
            "injected": 5000, 
            "tolerance": 200, 
            "unit": "RPM", 
            "type": "absolute", 
            "consigne": "5000 RPM ± 200 RPM"
        },
        
        # --- Température (Si Réfrigérée/Contrôlée) : Tolérance de +/- 10 °C ---
        "Précision Température (30 °C)": {
            "injected": 30, 
            "tolerance": 10, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "30 °C ± 10 °C"
        },
        "Précision Température (10 °C)": {
            "injected": 10, 
            "tolerance": 10, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "10 °C ± 10 °C"
        },
        "Précision Température (0 °C)": {
            "injected": 0, 
            "tolerance": 10, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "0 °C ± 10 °C"
        },
        "Précision Température (-10 °C)": {
            "injected": -10, 
            "tolerance": 10, 
            "unit": "°C", 
            "type": "absolute", 
            "consigne": "-10 °C ± 10 °C"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Fonctionnement des commandes (boutons, écran)": "Fonctionnel (Check)",
        "Vérification du temps de centrifugation programmé": "Conforme (Ok/Echoué/NA/NT)",
        "Contrôle mécanique (Roulements/freins, blocage capot)": "Conforme (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence de tous les accessoires (rotor, contrepoids, supports)",
        "Visibilité et lisibilité des inscriptions et des voyants",
        "État général (absence de rouille, couvercle, intérieur de la cuve)",
        "Fonctionnalité des roulements et freins des roulettes",
        "Fonctionnalité du blocage du capot/couvercle (sécurité essentielle)",
    ],
    
    # 4. Tests de Sécurité Électrique (Appareil de Classe I)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient : Non Applicable (Pas de contact patient)
        "Courant de fuite patient (parties appliquées)": {"limit": "NA", "unit": "mA", "type": "leakage"}, 
    }
},
    "Pompe à nutrition (Entérale)""": {
    # Informations Générales
    "NAME": "Pompe à nutrition (Entérale)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I)
    "FABRICANT_LIST": [
        "Fresenius Kabi",
        "Nestlé Health Science",
        "Moog",
        "Abbott/Hospira",
        "B. Braun",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Débit et Alarmes)
    "PERFORMANCE_CHECKS": {
        
        # Le contrôle de débit doit être effectué avec un ECME (testeur de débit)
        # Aucune valeur de consigne/seuil n'est fournie, donc c'est un check qualitatif/général ici.
        "Contrôle de Débit (avec ECME)": {
            "injected": "NA", # Valeur à définir selon le protocole de test (e.g., 100 mL/H)
            "tolerance": "NA", # Tolérance typique: 5-10%
            "unit": "mL/H", 
            "type": "NA", 
            "consigne": "Doit être conforme au débit programmé (utiliser ECME)"
        },
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Fonctionnement du clamp automatique": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement du détecteur de gouttes": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement sur batterie (capacité et charge)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Contrôle de débit général": "Conforme (Ok/Echoué/NA/NT)", # Résultat du test ECME (Réf 026)
        "Contrôle des alarmes (occlusion, air, fin de poche, sonores et visuelles)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil et du ventilateur d’alimentation",
        "Présence de tous les accessoires (câble secteur, potence/fixations)",
        "Visibilité et lisibilité des inscriptions et des voyants",
        "État général de l'appareil (boîtier, pompe péristaltique, capteurs)",
        "Vérification de la fonctionnalité des fixations (potence, rail)",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - contact patient via la ligne de perfusion)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Partie appliquée de Type BF)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Rampe chauffante (Réchauffeur néonatal)" : {
    # Informations Générales
    "NAME": "Rampe chauffante (Réchauffeur néonatal)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I/BF)
    "FABRICANT_LIST": [
        "Dräger",
        "GE Healthcare",
        "Atom Medical",
        "Ginevri",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Temps et Précision de Température)
    "PERFORMANCE_CHECKS": {
        
        # --- Mesures de Temps de Montée en Température (Qualitatif/Mesure) ---
        # Ces valeurs doivent être remplies et servent de base à l'évaluation Réf 028
        "Durée de passage de 23 à 30°C": {
            "injected": "NA", 
            "tolerance": "NA", 
            "unit": "min", 
            "type": "NA", 
            "consigne": "Mesure à remplir"
        },
        "Durée de passage de 30 à 36°C": {
            "injected": "NA", 
            "tolerance": "NA", 
            "unit": "min", 
            "type": "NA", 
            "consigne": "Mesure à remplir"
        },
        
        # Le contrôle final de la température (Réf 033) est un check qualitatif.
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Ok/Echoué)",
        "Fonctionnement de l'éclairage": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement sur batterie (capacité et charge)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Contrôle de la température autorégulée (<30 mn)": "Conforme (Ok/Echoué/NA/NT)", # Basé sur les mesures de temps
        "Contrôle de la température (enceinte, cutanée, couchage)": "Conforme (Ok/Echoué/NA/NT)",
        "Contrôle de la sécurité mécanique (roulettes, hauteur variable)": "Conforme (Ok/Echoué/NA/NT)",
        "Contrôle des alarmes (coupure alim., écart intensité, forte/faible intensité)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil (rampe, plan de couchage)",
        "Présence de tous les accessoires (câble secteur, capteur cutané, etc.)",
        "Visibilité et lisibilité des inscriptions et des voyants/alarmes",
        "État général de l'appareil (boîtier, résistance, surfaces chauffantes)",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - contact patient indirect)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Partie appliquée de Type BF - capteur cutané ou contact indirect par la chaleur)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Table d'opération" : {
    # Informations Générales
    "NAME": "Table d'opération",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I/BF)
    "FABRICANT_LIST": [
        "Maquet/Getinge",
        "Stryker",
        "Trumpf Medical",
        "Mizuho OSI",
        "Schmitz",
        "Autres"
    ],
    
    # 1. Tests de Performance (Principalement fonctionnels, pas de seuils numériques spécifiques fournis)
    "PERFORMANCE_CHECKS": {
        # Les contrôles de performance sont principalement mécaniques et qualitatifs pour une table d'opération.
        # En l'absence de seuils numériques pour les angles/hauteurs, cette section est laissée vide.
    },
    
    # 2. Tests Spécifiques (Contrôle de Fonctionnement et Sécurité Mécanique)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
        "Fonctionnement des mouvements (Hauteur, Trendelenburg, Latéralité, Segmentation)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement des fins de courses (Limites maximales/minimales atteintes)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement de la liaison au sol (Freinage/Blocage)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement des accessoires (Supports, sangles, têtes de lit/pieds)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement électrique (Commandes, Moteurs)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement du transfert (si modèle amovible/transférable)": "Conforme (Ok/Echoué/NA/NT)",
        "Contrôle global du fonctionnement (général)": "Conforme (Ok/Echoué/NA/NT)",
        "Contrôle des alarmes (Batterie faible, Surcharge, Mouvement bloqué, etc.)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Mécanique)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil",
        "Présence et intégrité de tous les accessoires/segments",
        "Visibilité et lisibilité des inscriptions et des voyants",
        "État général (absence de chocs, rouille, état des coussins/matelas)",
        "Vérification du fonctionnement des boutons et/ou de l’écran tactile",
        "Absence de fuite hydraulique et état des vérins/cylindres",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - contact patient)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Partie appliquée de Type BF - contact via le matelas/structure)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
   "Thermosoudeuse (Stérilisation)": {
        "NAME": "Thermosoudeuse (Stérilisation)",
        "TSE_REQUIRED": True, 
        "FABRICANT_LIST": ["Hawo", "Belimed", "Wipak", "Gandus", "Autres"],
        
        "PERFORMANCE_CHECKS": {
            "Contrôle Température de Soudage": {
                "injected": 180,        
                "tolerance": 5,         
                "unit": "°C", 
                "type": "measure",      
                "consigne": "Mesurer la température réelle avec une sonde externe."
            },
            "Vitesse de défilement": {
                "injected": 10, 
                "tolerance": 0.5, 
                "unit": "m/min", 
                "type": "measure",
                "consigne": "Vérifier la vitesse réelle de défilement."
            }
        },

        "SPECIFIC_CHECKS": {
            "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
            "Autotest de l'appareil au démarrage": "Fonctionnel (Check)",
            "Contrôle du roulement du tapis": "Conforme (Ok/Echoué/NA)",
            "Contrôle des alarmes": "Alarmes fonctionnelles (Ok/Echoué/NA)",
        },

        "VISUAL_CHECKS": [
            "État de propreté (barres, guides)",
            "État des accessoires (coupe-sachets)",
            "Lisibilité des voyants et écran",
            "État général (câble secteur, prise)",
        ],

        # Section Sécurité Électrique (Identique à l'Autoclave)
        "SECURITY_CHECKS": {
            "Résistance de terre": {
                "limit": 0.3, # < 0.3 Ω
                "unit": "Ω", 
                "type": "resistance"
            },
            "Résistance d'isolement": {
                "limit": 2.0, # > 2.0 MΩ
                "unit": "MΩ", 
                "type": "insulation"
            },
            "Courant de fuite châssis": {
                "limit": 0.5, # < 0.5 mA
                "unit": "mA", 
                "type": "leakage"
            }
        },
             
},
    "Ventilateur d'anesthésie": {
    # Informations Générales
    "NAME": "Ventilateur d'anesthésie",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I/BF)
    "FABRICANT_LIST": [
        "Dräger",
        "GE Healthcare",
        "Mindray",
        "Maquet/Getinge",
        "Datex Ohmeda",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Précision des paramètres)
    "PERFORMANCE_CHECKS": {
        
        # --- Contrôle en Volume Contrôlé (VC) ---
        "VC - FiO2 (%)": {
            "injected": 21, "tolerance": 2.1, "unit": "%", "type": "Mesure", 
            "consigne": 21.0, "seuil_bas": 18.9, "seuil_haut": 23.1
        },
        "VC - Volume Minute (L/min)": {
            "injected": 5.0, "tolerance": 0.5, "unit": "L/min", "type": "Mesure", 
            "consigne": 5.0, "seuil_bas": 4.5, "seuil_haut": 5.5
        },
        "VC - Volume Courant (Vt) (mL)": {
            "injected": 500, "tolerance": 50, "unit": "mL", "type": "Mesure", 
            "consigne": 500, "seuil_bas": 450, "seuil_haut": 550
        },
        "VC - Pression Max (Pmax) (cmH2O)": {
            "injected": 40, "tolerance": 4, "unit": "cmH2O", "type": "Mesure", 
            "consigne": 40.0, "seuil_bas": 36.0, "seuil_haut": 44.0
        },
        "VC - Fréquence (Freq) (c/min)": {
            "injected": 10, "tolerance": 1, "unit": "c/min", "type": "Mesure", 
            "consigne": 10.0, "seuil_bas": 9.0, "seuil_haut": 11.0
        },
        "VC - Rapport I/E": {
            "injected": 0.5, "tolerance": 0.05, "unit": "Ratio", "type": "Mesure", 
            "consigne": 0.5, "seuil_bas": 0.45, "seuil_haut": 0.55
        },
        "VC - PEEP (cmH2O)": {
            "injected": 4, "tolerance": 0.4, "unit": "cmH2O", "type": "Mesure", 
            "consigne": 4.0, "seuil_bas": 3.6, "seuil_haut": 4.4
        },

        # --- Contrôle en Pression Contrôlée (PC) ---
        "PC - FiO2 (%)": {
            "injected": 21, "tolerance": 2.1, "unit": "%", "type": "Mesure", 
            "consigne": 21.0, "seuil_bas": 18.9, "seuil_haut": 23.1
        },
        "PC - Pression Inspiratoire (cmH20)": {
            # Note: Consigne basée sur la Réf 042 (30) malgré le réglage à 20 dans Réf 040.
            "injected": 30, "tolerance": 3, "unit": "cmH2O", "type": "Mesure", 
            "consigne": 30.0, "seuil_bas": 27.0, "seuil_haut": 33.0
        },
        "PC - Fréquence (Freq) (c/min)": {
            "injected": 10, "tolerance": 1, "unit": "c/min", "type": "Mesure", 
            "consigne": 10.0, "seuil_bas": 9.0, "seuil_haut": 11.0
        },
        "PC - Rapport I/E": {
            "injected": 0.5, "tolerance": 0.05, "unit": "Ratio", "type": "Mesure", 
            "consigne": 0.5, "seuil_bas": 0.45, "seuil_haut": 0.55
        },
        "PC - PEEP (cmH2O)": {
            "injected": 4, "tolerance": 0.4, "unit": "cmH2O", "type": "Mesure", 
            "consigne": 4.0, "seuil_bas": 3.6, "seuil_haut": 4.4
        },
    },
    
    # 2. Tests Spécifiques (Fonctionnel & Sécurité gazeuse)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Fonctionnement des rotamètres (Air et O2)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement sur batterie (capacité et charge)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Vérification des tests de fuite interne (appareil)": "Conforme (Ok/Echoué/NA/NT)",
        "Vérification des tests de fuite du ballon manuel": "Conforme (Ok/Echoué/NA/NT)",
        "Étalonnage de la cellule à O2 à 21 %": "Conforme (Ok/Echoué/NA/NT)",
        "Étalonnage de la cellule à O2 à 100 % (si gaz disponible)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement du Trigger (VC/PC)": "Fonctionnel (Ok/Echoué)",
        "Contrôle global des alarmes (y compris défauts d'alimentation gaz/électrique)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Accessoires)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil",
        "Présence et intégrité de tous les accessoires (capteurs, tuyaux, filtres, chaux sodée, etc.)",
        "Visibilité et lisibilité des inscriptions et des voyants/alarmes",
        "État général (boîtier, bras, roulettes, freins)",
        "Vérification de la fixation des tuyaux, des circuits et du filtre/piston",
        "Fonctionnement des boutons et/ou de l’écran tactile",
    ],
    
    # 4. Tests de Sécurité Électrique (Type BF - contact patient)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient (Partie appliquée de Type BF)
        "Courant de fuite patient (parties appliquées)": {"limit": "< 100 μA (0.100 mA)", "unit": "mA", "type": "leakage"}, 
    }
},
    "Concentrateur d'oxygène" : {
    # Informations Générales
    "NAME": "Concentrateur d'oxygène",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I/BF)
    "FABRICANT_LIST": [
        "Invacare",
        "Philips Respironics",
        "AirSep/Caire",
        "Devilbiss",
        "Yuwell",
        "Autres"
    ],
    
    # 1. Tests de Performance (Mesures critiques : Débit et Concentration d'O2)
    "PERFORMANCE_CHECKS": {
        
        # --- Contrôle du Débit (LPM) ---
        "Débit à 2 LPM (L/min)": {
            "injected": 2.0, "tolerance": 0.2, "unit": "L/min", "type": "Mesure", 
            "consigne": 2.0, "seuil_bas": 1.8, "seuil_haut": 2.2
        },
        "Débit à 5 LPM (L/min)": {
            "injected": 5.0, "tolerance": 0.5, "unit": "L/min", "type": "Mesure", 
            "consigne": 5.0, "seuil_bas": 4.5, "seuil_haut": 5.5
        },
        "Débit à 10 LPM (L/min) (si applicable)": {
            "injected": 10.0, "tolerance": 1.0, "unit": "L/min", "type": "Mesure", 
            "consigne": 10.0, "seuil_bas": 9.0, "seuil_haut": 11.0
        },
        
        # --- Contrôle de la Concentration (FiO2) ---
        "Concentration O2 à 5 LPM (%)": {
            # Note: Le seuil haut n'est pas spécifié, seulement le seuil bas (90%). La cible est 93%.
            "injected": 93, "tolerance": 3, "unit": "%", "type": "Mesure", 
            "consigne": 93.0, "seuil_bas": 90.0, "seuil_haut": "NA"
        },
        
        # --- Contrôle de la Pression (Qualitatif/Documentaire) ---
        "Pression de sortie à 5 LPM": {
            "injected": "NA", "tolerance": "NA", "unit": "Mesure", "type": "Mesure", 
            "consigne": "Mesure conforme à la documentation technique (±10%)"
        },
    },
    
    # 2. Tests Spécifiques (Fonctionnel & Alarmes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Fonctionnement du rotamètre (stabilité et fluidité)": "Conforme (Ok/Echoué/NA/NT)",
        "Fonctionnement sur batterie (si applicable)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Alarme de coupure d’alimentation (batterie système d'alarme)": "Alarme fonctionnelle (Ok/Echoué/NA/NT)",
        "Alarme de défaut technique (obturation de la sortie O2)": "Alarme fonctionnelle (Ok/Echoué/NA/NT)",
        "Contrôle mécanique (roulettes, freins, fixations)": "Conforme (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Filtres)
    "VISUAL_CHECKS": [
        "État de propreté de l'appareil",
        "Présence et intégrité de tous les accessoires (câble secteur, humidificateur, olive de sortie, etc.)",
        "Visibilité et lisibilité des inscriptions et des voyants/alarmes",
        "État des filtres (poussière, longue durée, antibactérien) et maintenance à effectuer (nettoyage/remplacement)",
        "Vérification de la fixation des tuyaux, capots et humidificateur",
        "Fonctionnement des boutons et/ou de l’écran tactile",
    ],
    
    # 4. Tests de Sécurité Électrique (Appareil de Classe I)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard pour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient : Non Applicable (Pas de contact patient direct par l'appareil)
        "Courant de fuite patient (parties appliquées)": {"limit": "NA", "unit": "mA", "type": "leakage"}, 
    }
},
    "Reprographe film RX (Imageur sec)" : {
    # Informations Générales
    "NAME": "Reprographe film RX (Imageur sec)",
    "TSE_REQUIRED": True, # Test de Sécurité Électrique requis (appareil de Classe I)
    "FABRICANT_LIST": [
        "Agfa",
        "Kodak / Carestream",
        "FujiFilm",
        "Drystar",
        "Autres"
    ],
    
    # 1. Tests de Performance (Qualité de l'Image et Mécanique)
    "PERFORMANCE_CHECKS": {
        # Les tests de performance se concentrent sur la qualité visuelle du film
        "Nettoyage de la tête d'impression": {
            "injected": "NA", "tolerance": "NA", "unit": "Entretien", "type": "Ok/Echoué", 
            "consigne": "Réaliser le nettoyage si nécessaire ou selon les recommandations"
        },
        "Nettoyage des rouleaux d'entraînement": {
            "injected": "NA", "tolerance": "NA", "unit": "Entretien", "type": "Ok/Echoué", 
            "consigne": "Réaliser le nettoyage pour éviter les artefacts d'impression"
        },
        "Qualité d'impression (Bac 1)": {
            "injected": "NA", "tolerance": "NA", "unit": "Contrôle Qualité", "type": "Ok/Echoué", 
            "consigne": "Vérifier la lisibilité des 3 groupes de 5 points dans chaque ovale de la mire de contrôle (Réf 026)"
        },
        "Qualité d'impression (Bac 2)": {
            "injected": "NA", "tolerance": "NA", "unit": "Contrôle Qualité", "type": "Ok/Echoué", 
            "consigne": "Vérifier la lisibilité des 3 groupes de 5 points dans chaque ovale de la mire de contrôle (Réf 027)"
        },
        "Qualité d'impression (Bac 3) (si applicable)": {
            "injected": "NA", "tolerance": "NA", "unit": "Contrôle Qualité", "type": "Ok/Echoué", 
            "consigne": "Vérifier la lisibilité des 3 groupes de 5 points dans chaque ovale de la mire de contrôle (Réf 028)"
        },
        "Qualité d'impression (Bac 4) (si applicable)": {
            "injected": "NA", "tolerance": "NA", "unit": "Contrôle Qualité", "type": "Ok/Echoué", 
            "consigne": "Vérifier la lisibilité des 3 groupes de 5 points dans chaque ovale de la mire de contrôle (Réf 029)"
        },
    },
    
    # 2. Tests Spécifiques (Fonctionnel & Automatismes)
    "SPECIFIC_CHECKS": {
        "L'équipement ne fait pas l'objet d'une matériovigilance": "Vérification dans les registres (Oui/Non)",
        "Autotest de l'appareil (au démarrage)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Fonctionnement sur batterie (si applicable)": "Fonctionnel (Ok/Echoué/NA/NT)",
        "Reconnaissance automatique des films (formats et types)": "Conforme (Ok/Echoué/NA/NT)",
        "Contrôle des alarmes (manque de film, erreur d'impression, température)": "Alarmes fonctionnelles (Ok/Echoué/NA/NT)",
    },
    
    # 3. Contrôles Visuels (État Général et Propreté)
    "VISUAL_CHECKS": [
        "Nombre de bacs d'impression disponibles",
        "Nettoyage des orifices de refroidissement",
        "État de propreté de l'appareil",
        "Présence et intégrité de tous les accessoires (câble secteur, câble réseau)",
        "Visibilité et lisibilité des inscriptions",
        "État général (absence de choc, fonctionnement des roulettes/freins si mobile)",
        "Fonctionnement des boutons et/ou de l’écran tactile",
    ],
    
    # 4. Tests de Sécurité Électrique (Appareil de Classe I)
    "SECURITY_CHECKS": {
        # Résistance de terre (standard pour appareil de Classe I)
        "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        # Courant de fuite du châssis (standard lpour appareil de Classe I)
        "Courant de fuite châssis": {"limit": "< 500 μA (0.500 mA)", "unit": "mA", "type": "leakage"},
        # Courant de fuite patient : Non Applicable (Pas de contact patient)
        "Courant de fuite patient (parties appliquées)": {"limit": "NA", "unit": "mA", "type": "leakage"}, 
    }
},
    "Lève-malade": {
        "NAME": "Lève-malade / Verticalisateur",
        "TSE_REQUIRED": True, # Correction : Toujours True pour activer l'affichage
        "FABRICANT_LIST": ["Arjo", "Invacare", "Hill-Rom", "Autre"],
        "PERFORMANCE_CHECKS": {
            "Batterie (Tension charge)": {"injected": 24, "tolerance": 2, "unit": "V", "type": "range"},
            "Vitesse de levée": {"injected": 10, "tolerance": 5, "unit": "s", "type": "absolute"},
        },
        "SPECIFIC_CHECKS": {
            "Descente d'urgence mécanique": "Opérationnelle",
            "Arrêt d'urgence électrique": "Opérationnel",
            "Limiteur de charge": "Vérifié"
        },
        "VISUAL_CHECKS": [
            "État des soudures du châssis", 
            "État des sangles/harnais", 
            "Fixation des roues et freins"
        ],
        "SECURITY_CHECKS": {
            # Résistance de terre (standard pour appareil de Classe I)
            "Résistance de terre": {"limit": "< 0.300 Ω", "unit": "Ω", "type": "resistance"}, 
        
        }
    },
    "Stérilisateur": {
        "NAME": "Stérilisateur / Autoclave",
        "TSE_REQUIRED": False,
        "FABRICANT_LIST": ["Melag", "Tuttnauer", "Getinge", "Autre"],
        "PERFORMANCE_CHECKS": {
            "Température (121 °C)": {"injected": 121, "tolerance": 1, "unit": "°C", "type": "range"},
            "Pression (2 bar)": {"injected": 2, "tolerance": 0.1, "unit": "bar", "type": "range"},
        },
        "SPECIFIC_CHECKS": {
            "Cycle B (Test de pénétration)": "Résultat du test réussi.",
            "Test d'étanchéité à vide": "Taux de fuite dans les limites du fabricant.",
        },
        "VISUAL_CHECKS": ["Vérifier le joint de porte et l'état de la cuve", "Vérifier les filtres à air et à eau"],
        "SECURITY_CHECKS": {}
    },
}

DEVICE_TYPES = ["--- Sélectionner un Type d'Équipement ---"] + ["Autre dispositif"] + list(QC_CONFIGS.keys())

SAFETY_LIMITS_COMMON = {
    "resistance": {"limit": 0.3, "unit": "Ω"},
    "fuite_patient": {"limit": 0.100, "unit": "mA"}, # 10 µA
    "fuite_chassis": {"limit": 0.500, "unit": "mA"}, # 500 µA
}


# --------------------------------------------------------------------------------------
# --- FONCTIONS DE VALIDATION (Laisser inchangé) ---------------------------------------
# --------------------------------------------------------------------------------------

def valider_performance(injectee: float, lue: float, tolerance: float, unit: str) -> tuple[str, bool]:
    """Vérifie si la valeur lue est dans la tolérance."""
    if lue is None or np.isnan(lue) or injectee is None or tolerance is None:
        return "N/A (Saisie incomplète/non numérique)", False

    ecart = abs(lue - injectee)
    conforme = ecart <= tolerance

    statut = (f"Conforme (Écart: {ecart:.2f} {unit})" if conforme
              else f"NON CONFORME (Écart: {ecart:.2f} > {tolerance:.2f} {unit})")

    return statut, conforme

def valider_securite(mesure: float, limite_ma: Any, type_test: str) -> tuple[str, bool, str, str]:
    """Vérifie si les mesures de sécurité sont dans les limites."""
    
    # 1. Determine the numeric limit and the unit string
    unit = "Ω" if type_test == "resistance" else "mA"
    
    # Use global default if it's a resistance test 
    if type_test == "resistance":
        limite_val = 0.3 # Default for resistance
    else:
        limite_val = limite_ma 

    # 2. Safety check: Convert to float to avoid the TypeError
    try:
        # If limite_val is a string like "< 0.500", extract the numeric part
        if isinstance(limite_val, str):
            import re
            # Extract numbers/decimals from string
            numeric_parts = re.findall(r"[-+]?\d*\.\d+|\d+", limite_val)
            limite_val = float(numeric_parts[0]) if numeric_parts else 0.5
        else:
            limite_val = float(limite_val)
    except (ValueError, TypeError, IndexError):
        # Fallback defaults
        limite_val = 0.3 if type_test == "resistance" else 0.5

    # 3. Perform the calculation
    conforme = mesure <= limite_val

    # 4. Define the missing strings required for the return statement
    lim_str = f"≤ {limite_val:.3f} {unit}"
    mesure_str = f"{mesure:.3f} {unit}"
    
    if conforme:
        statut_detail = f"Conforme ({mesure_str} {lim_str})"
    else:
        statut_detail = f"NON CONFORME ({mesure_str} > {limite_val:.3f} {unit})"

    # 5. Return (Fixed Indentation)
    return statut_detail, conforme, lim_str, mesure_str
# --------------------------------------------------------------------------------------
# --- FONCTIONS DE GÉNÉRATION PDF (REPORTLAB) - MISE À JOUR POUR LOGOS ---
# --------------------------------------------------------------------------------------

def first_page_header_logo(canvas, doc, config_name):
    """Dessine l'en-tête pour la PREMIÈRE PAGE (MODIFIÉ POUR LOGO)."""
    canvas.saveState()
    page_width = A4[0]
    margin = doc.leftMargin
    top_y = A4[1] - margin + 20 
    left_x = margin
    
    # --- GAUCHE: Logo AIMA --- (Cette partie reste inchangée pour garder le logo)
    logo_height = 75
    logo_width = 120  # Ajustez si nécessaire pour respecter les proportions
    try:
        # Tente d'utiliser le chemin du logo AIMA (assurez-vous que le chemin est correct)
        canvas.drawImage(AIMA_LOGO_PATH, left_x, top_y - logo_height, 
                         width=logo_width, height=logo_height, 
                         mask='auto')
        text_start_y = top_y - logo_height - 15 
        
    except:
        # Fallback au texte si le logo n'est pas trouvé ou si le chemin est invalide
        canvas.setFont('Helvetica-Bold', 12)
        canvas.drawString(left_x, top_y - 30, "Allons Imaginer un")
        canvas.drawString(left_x, top_y - 45, "Monde d'Amitiés")
        text_start_y = top_y - 40 # Position du titre du rapport

    # ----------------------------------------------------------------
    # DROITE: Titre du Rapport
    # ----------------------------------------------------------------
    right_x_end = page_width - margin
    
    # 1. ATTESTATION DE CONTROLE QUALITÉ ET DES PERFORMANCES
    title_text = "ATTESTATION DE CONTROLE QUALITÉ ET DES PERFORMANCES"
    title_font_size = 12
    canvas.setFont('Helvetica-Bold', title_font_size)
    text_width_main = canvas.stringWidth(title_text, 'Helvetica-Bold', title_font_size)
    canvas.drawString(right_x_end - text_width_main, top_y - 20, title_text)
    
    # 2. DISPOSITIF MEDICAL (DM): ...
    device_text = f"DISPOSITIF MEDICAL (DM): {config_name}"
    device_font_size = 9
    canvas.setFont('Helvetica-Bold', device_font_size)
    text_width_dev = canvas.stringWidth(device_text, 'Helvetica-Bold', device_font_size)
    y_device_title = top_y - 38
    canvas.drawString(right_x_end - text_width_dev, y_device_title, device_text)
    
    # LIGNE SUPPRIMÉE:
    # Les lignes de code qui dessinaient la ligne noire horizontale sont maintenant omises.
    
    canvas.restoreState()

def first_page_footer_logos(canvas, doc, config_name):
    """Dessine le bloc de logos au bas de la PREMIÈRE PAGE (NOUVEAU)."""
    canvas.saveState()
    margin = doc.leftMargin
    bottom_y = doc.bottomMargin / 2 # Laisse un peu de place
    
    # ----------------------------------------------------------------
    # NOUVEAU: BLOC DE LOGOS EN BAS À GAUCHE
    # ----------------------------------------------------------------
    try:
        # Tente d'utiliser le chemin du bloc de logos
        logos_block = ImageReader(BOTTOM_LOGOS_PATH) 
        
        # Largeur et Hauteur maximales pour le bloc (ajustez ces valeurs au besoin)
        max_width = 2200
        max_height = 100
        
        logo_width = max_width 
        logo_height = max_width * logos_block.getSize()[1] / logos_block.getSize()[0]
        
        # Si la hauteur calculée dépasse la limite, redimensionner par la hauteur
        if logo_height > max_height:
             logo_height = max_height
             logo_width = max_height * logos_block.getSize()[0] / logos_block.getSize()[1]

        # Positionnement: left_x, doc.bottomMargin
        canvas.drawImage(logos_block, margin, doc.bottomMargin, width=logo_width, height=logo_height)
        
    except:
        # Fallback ou ignorer si le logo n'est pas trouvé
        pass 
        
    canvas.restoreState()


def later_page_header_logo(canvas, doc, config_name):
    """Dessine un numéro de page en bas."""
    canvas.saveState()
    page_num = canvas.getPageNumber()
    text = f"Page {page_num}"
    canvas.setFont('Helvetica', 8)
    # Positionné en bas à droite
    canvas.drawString(A4[0] - doc.rightMargin - 30, doc.bottomMargin / 2, text)
    canvas.restoreState()


def generate_pdf_report(data: Dict[str, Any], config: Dict[str, Any]):
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=50,
                            bottomMargin=30)
    doc.config_name = config.get('NAME', 'RAPPORT DE CONTROLE QUALITÉ')

    # Styles (Modifié pour réduire l'espacement vertical des titres)
    styles = getSampleStyleSheet()
    Story: List[Any] = []
    
    # Style des titres de section principaux (I., II., III., etc.)
    style_heading2 = styles['Heading2']
    style_heading2.spaceBefore = 6   # RÉDUIT l'espace avant le titre (était par défaut 12-18)
    style_heading2.spaceAfter = 3    # RÉDUIT l'espace après le titre (était 0)
    style_heading2.fontSize = 10    
    
    # Style des sous-titres (II.A, II.B, etc.)
    style_heading3 = styles['Heading3']
    style_heading3.spaceBefore = 3   # RÉDUIT l'espace avant le sous-titre (était par défaut 6)
    style_heading3.spaceAfter = 3    # RÉDUIT l'espace après le sous-titre (était 1)
    style_heading3.fontSize = 9     
    
    style_normal = styles['Normal']
    style_normal.fontSize = 8   
    style_bold = styles['Code']     
    style_bold.fontSize = 8
    style_id_bold = style_bold
    style_id_normal = style_normal
    
    # Style pour les clauses (réduit l'espace autour du texte de la clause, était 5/5)
    style_clause_body = ParagraphStyle('ClauseBody', parent=style_normal, spaceAfter=2, spaceBefore=2, fontSize=7)
    
    base_table_commands = [
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'), 
        ('FONTSIZE', (0,0), (-1,-1), 8), 
        ('LEFTPADDING', (0,0), (-1,-1), 1),
        ('RIGHTPADDING', (0,0), (-1,-1), 1),
        ('TOPPADDING', (0,0), (-1,-1), 2),    
        ('BOTTOMPADDING', (0,0), (-1,-1), 2), 
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F2F2F2')),
        ('ALIGN', (0,0), (0,-1), 'LEFT'),  
        ('ALIGN', (1,1), (1,-1), 'CENTER'),
    ]
    
    Story.append(Spacer(1, 40)) # Augmenter l'espace initial pour laisser place au logo
    # ... (Ajout du contenu inchangé)
    # --- I. Identification ---
    Story.append(Paragraph("<b>I. Identification de l'Équipement</b>", style_heading2))

    # SÉCURISATION DES DONNÉES
    fabricant_raw = data.get('fabricant', 'N/A')
    fabricant_display = str(fabricant_raw)
    if fabricant_display == 'Autre':
        fabricant_display = str(data.get('autre_fabricant', 'Non Spécifié'))
    modele = str(data.get('modele', 'N/A'))
    sn = str(data.get('sn', 'N/A'))
    aima_sn = str(data.get('aima_sn', 'N/A'))
    date_str = str(data.get('date', 'N/A')) 
    lieu = str(data.get('lieu', 'N/A'))
    technicien = str(data.get('technicien', 'N/A'))
    ecme = str(data.get('ecme', 'N/A'))
    
    # Structure compacte 4x4
    id_data = [
        # Ligne d'en-tête fusionnée pour obtenir la bande grise uniforme
        [Paragraph("<b>Informations&nbsp;Principales&nbsp;de&nbsp;l'Équipement</b>", style_bold), '', '', ''], 
        
        # Lignes de données (Label A / Value A / Label B / Value B)
        [
            Paragraph("<b>Fabriquant:</b>", style_id_bold), Paragraph(fabricant_display, style_normal), 
            Paragraph("<b>Modèle&nbsp;/&nbsp;Référence:</b>", style_id_bold), Paragraph(modele, style_normal)
        ],
        [
            Paragraph("<b>N°&nbsp;serie:</b>", style_id_bold), Paragraph(sn, style_normal), 
            Paragraph("<b>N°&nbsp;série&nbsp;AIMA&nbsp;(GDR):</b>", style_id_bold), Paragraph(aima_sn, style_normal)
        ],
        [
            Paragraph("<b>Date&nbsp;de&nbsp;controle:</b>", style_id_bold), Paragraph(date_str, style_normal), 
            Paragraph("<b>Lieu&nbsp;de&nbsp;controle:</b>", style_id_bold), Paragraph(lieu, style_normal)
        ],
        [
            Paragraph("<b>Nom&nbsp;de&nbsp;l'intervenant&nbsp;technique:</b>", style_id_bold), Paragraph(technicien, style_normal), 
            Paragraph("<b>ECME&nbsp;utilisé:</b>", style_id_bold), Paragraph(ecme, style_normal)
        ],
    ]

    # Style spécifique pour la fusion de la première ligne
    custom_id_style_commands = list(base_table_commands) 
    custom_id_style_commands.append(('SPAN', (0,0), (-1,0))) # Fusionne les cellules de la première ligne
    custom_id_style_commands.append(('ALIGN', (0,0), (-1,0), 'CENTER')) 
    custom_id_style_commands.append(('ALIGN', (0, 1), (0, -1), 'LEFT')) 
    custom_id_style_commands.append(('ALIGN', (2, 1), (2, -1), 'LEFT')) 
    
    custom_id_style = TableStyle(custom_id_style_commands)

    id_table = Table(id_data, colWidths=[140, 140, 140, 140]) 
    id_table.setStyle(custom_id_style)
    Story.append(id_table)
    Story.append(Spacer(1, 2)) 


    # --- II. Vérifications Générales et Visuelles ---
    Story.append(Paragraph("<b>II. Vérifications Générales et Visuelles</b>", style_heading2))

    # II.A Matériovigilance / Alertes de Sécurité
    Story.append(Paragraph("<b>II.A Matériovigilance / Alertes de Sécurité</b>", style_heading3))
    matvig_status_raw = data.get('materiovigilance')
    matvig_status = '❌ NON' if matvig_status_raw == 'non' or matvig_status_raw is None else '✅ OUI (Alerte confirmée)'
    matvig_comment = str(data.get('commentaire_matvig', 'N/A'))
    
    table_matvig_data = [
        [
            Paragraph("<b>Contrôle</b>", style_bold),
            Paragraph("<b>Statut&nbsp;</b>", style_bold), 
            Paragraph("<b>Commentaire / Observation</b>", style_bold),
        ],
        [
            Paragraph("L'équipement **n'a pas** fait l'objet d'une matériovigilance / alerte :", style_normal),
            Paragraph(matvig_status, style_normal),
            Paragraph(matvig_comment, style_normal),
        ]
    ]
    table_matvig_col_widths = [340, 70, 150]

    table_matvig = Table(table_matvig_data, colWidths=table_matvig_col_widths)
    style_matvig = list(base_table_commands)
    style_matvig.append(('ALIGN', (1, 1), (1, -1), 'CENTER'))
    table_matvig.setStyle(TableStyle(style_matvig))
    Story.append(table_matvig)
    Story.append(Spacer(1, 2)) 


    # II.B Contrôles Visuels / Intégrité Physique
    Story.append(Paragraph("<b>II.B Contrôles Visuels / Intégrité Physique</b>", style_heading3))

    table_II_data = [
        [
            Paragraph("<b>Étape de Contrôle</b>", style_bold),
            Paragraph("<b>OK&nbsp;</b>", style_bold), 
            Paragraph("<b>Commentaire / Observation</b>", style_bold),
        ]
    ]
    table_II_col_widths = [350, 60, 150]
    
    # Visuels/Généraux
    for i, check in enumerate(config.get('VISUAL_CHECKS', [])):
        key = f'visuel_check_{i}'
        val = data.get(key, 'N/A')
        
        if isinstance(val, bool):
             response_text = '✅' if val else '❌' 
        else:
            response_text = 'N/A'
        
        check_key = check.replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "").lower()
        comment_key = f"comment_visuel_{check_key}" 
        comment = str(data.get(comment_key, ''))
        
        table_II_data.append([
            Paragraph(f"• {check}", style_normal),
            Paragraph(response_text, style_normal),
            Paragraph(comment, style_normal),
        ])

    table_II = Table(table_II_data, colWidths=table_II_col_widths)
    style_ii = list(base_table_commands)
    style_ii.append(('ALIGN', (1, 1), (1, -1), 'CENTER'))
    table_II.setStyle(TableStyle(style_ii))
    Story.append(table_II)
    Story.append(Spacer(1, 2)) 


    # --- III. Tests de Performance ---
    Story.append(Paragraph("<b>III. Tests de Performance et Exactitude</b>", style_heading2))
    
    table_III_data = [
        [
            Paragraph("<b>Test&nbsp;Paramètre</b>", style_bold),
            Paragraph("<b>Statut&nbsp;</b>", style_bold), 
            Paragraph("<b>Valeur&nbsp;Injectée/Tolérance</b>", style_bold),
            Paragraph("<b>Valeur&nbsp;Lue/Observation</b>", style_bold),
        ]
    ]
    table_III_col_widths = [150, 70, 170, 170] 
    
    for check_name_key, result in data.get('performance_results', {}).items():
        
        statut = result.get('statut', 'N/A')
        
        if 'valeur' in result and 'injected' in result:
            display_name = result.get('check_name', check_name_key)
            lue = result['valeur']
            try:
                lue_formatted = f"{lue:.2f}".rstrip('0').rstrip('.') if lue is not None and not np.isnan(lue) else 'N/A'
            except (TypeError, ValueError):
                lue_formatted = str(lue or 'N/A')
            
            seuil_text = f"{result['injected']} {result['unit']} (±{result['tolerance']} {result['unit']})"
            statut_display = '✅ OK' if 'Conforme' in statut else ('❌ NON OK' if 'NON CONFORME' in statut else 'N/A')
            
            obs_text = f"Lue: {lue_formatted} {result.get('unit', '')}. "
            comment = result.get('comment', '')
            if comment:
                 obs_text += f"Obs: {comment}"
            else:
                 obs_text += f"Obs: N/A"

            table_III_data.append([
                Paragraph(display_name, style_normal),
                Paragraph(statut_display, style_normal),
                Paragraph(seuil_text, style_normal),
                Paragraph(obs_text, style_normal),
            ])

    table_III = Table(table_III_data, colWidths=table_III_col_widths)
    style_iii = list(base_table_commands)
    style_iii.append(('ALIGN', (1, 1), (1, -1), 'CENTER'))
    table_III.setStyle(TableStyle(style_iii))
    Story.append(table_III)
    Story.append(Spacer(1, 2)) 


    # --- IV. Tests Spécifiques ---
    Story.append(Paragraph("<b>IV. Tests Spécifiques</b>", style_heading2))
    
    table_IV_data = [
        [
            Paragraph("<b>Test&nbsp;Spécifique</b>", style_bold),
            Paragraph("<b>Résultat Attendu</b>", style_bold),
            Paragraph("<b>Statut&nbsp;</b>", style_bold), 
            Paragraph("<b>Observation</b>", style_bold),
        ]
    ]
    table_IV_col_widths = [150, 180, 70, 160] 
    
    all_specific_checks = {} 
    # Ajouter les tests prédéfinis
    for name, expected in config.get('SPECIFIC_CHECKS', {}).items():
        all_specific_checks[name] = expected
    
    # Ajouter les tests manuels à la liste
    for key in data.get('manual_spec_keys', []):
        all_specific_checks[key] = { 
            'check_name': data.get(f"{key}_name", f"Test Manuel {key.split('_')[-1]}"),
            'expected_result': data.get(f"{key}_expected", "N/A (Saisie Manuelle)")
        }

    for check_name_key, expected_result_or_data in all_specific_checks.items():
        if isinstance(expected_result_or_data, str):
            check_name = check_name_key
            expected_result = expected_result_or_data
            result_passed = data.get('specific_results', {}).get(check_name, 'N/A')
            comment_key = check_name 
        else:
            check_name = expected_result_or_data['check_name']
            expected_result = expected_result_or_data['expected_result']
            result_passed = data.get('specific_results', {}).get(check_name_key, 'N/A')
            comment_key = check_name_key 

        if isinstance(result_passed, bool):
             response_text = '✅ OK' if result_passed else '❌ NON OK'
        else:
             response_text = 'N/A' # Si le test n'a pas été coché/décoché

        comment = str(data.get(f'comment_spec_{comment_key}', ''))
        
        table_IV_data.append([
            Paragraph(check_name, style_normal),
            Paragraph(expected_result, style_normal),
            Paragraph(response_text, style_normal),
            Paragraph(comment, style_normal),
        ])

    table_IV = Table(table_IV_data, colWidths=table_IV_col_widths)
    style_iv = list(base_table_commands)
    style_iv.append(('ALIGN', (2, 1), (2, -1), 'CENTER'))
    table_IV.setStyle(TableStyle(style_iv))
    Story.append(table_IV)
    Story.append(Spacer(1, 2)) 

    # --- V. Tests de Sécurité Électrique ---
    Story.append(Paragraph("<b>V. Tests de Sécurité Électrique</b>", style_heading2))
    
    tse_required = config.get("TSE_REQUIRED", True)

    if not tse_required:
        table_V_data = [
            [
                Paragraph("<b>Test Électrique</b>", style_bold),
                Paragraph("<b>Statut</b>", style_bold),
                Paragraph("<b>Commentaire</b>", style_bold),
            ],
            [
                Paragraph("Tests de Sécurité Électrique (IEC 62353)", style_normal),
                Paragraph("N/A", style_normal),
                Paragraph(f"Ce dispositif ({config['NAME']}) n'est pas soumis à l'IEC 62353 et relève d'une autre norme de sécurité (ex : IEC 62353 pour Labo/Stérilisation, ou Classe II/batterie sans terre).", style_normal),
            ]
        ]
        table_V = Table(table_V_data, colWidths=[300, 70, 190])
        style_v = list(base_table_commands)
        style_v.append(('ALIGN', (1, 1), (1, -1), 'CENTER'))
        table_V.setStyle(TableStyle(style_v))
        Story.append(table_V)

    else:
        table_V_data = [
            [
                Paragraph("<b>Test&nbsp;électrique</b>", style_bold),
                Paragraph("<b>Limite</b>", style_bold),
                Paragraph("<b>Mesure</b>", style_bold),
                Paragraph("<b>Statut&nbsp;</b>", style_bold), 
            ]
        ]
        table_V_col_widths = [150, 150, 150, 100] 
        
        secu_items_config = [
            'Résistance de terre',
            'Courant de fuite patient',
            'Courant de fuite châssis',
        ]
        
        for check_name in secu_items_config:
            if check_name not in config['SECURITY_CHECKS']:
                 continue
            
            result = data.get('security_results', {}).get(check_name, {})
            
            statut_display = result.get('statut', 'N/A')
            lim_display = result.get('lim_str', 'N/A')
            mesure_display = result.get('mesure_str', 'N/A')
            
            if statut_display == 'N/A':
                 statut_color = 'N/A'
            else:
                 statut_color = '✅ OK' if 'Conforme' in statut_display else '❌ NON OK'
            
            table_V_data.append([
                Paragraph(check_name.replace(" ", "&nbsp;"), style_normal), 
                Paragraph(lim_display, style_normal),
                Paragraph(mesure_display, style_normal),
                Paragraph(statut_color, style_normal),
            ])

        table_V = Table(table_V_data, colWidths=table_V_col_widths)
        style_v = list(base_table_commands)
        style_v.append(('ALIGN', (3, 1), (3, -1), 'CENTER'))
        table_V.setStyle(TableStyle(style_v))
        Story.append(table_V)

    Story.append(Spacer(1, 2)) 


    # --- VI. CONCLUSION ET VALIDATION ---
    Story.append(Paragraph("<b>VI. CONCLUSION ET VALIDATION</b>", style_heading2))
    
    global_conformity = data.get('global_conformity', 'INCONNU')
    
    conclusion_data = [
        [
            Paragraph("<b>Statut Final de l'Équipement:</b>", style_bold),
            Paragraph(global_conformity, styles['Heading2']),
        ],
        [
            Paragraph("<b>Commentaire Final:</b>", style_bold),
            Paragraph(data.get('final_comment', 'N/A (Rapport généré avant complétion)'), style_normal), 
        ]
    ]
    conclusion_table = Table(conclusion_data, colWidths=[150, 410])
    conclusion_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('ALIGN', (1,0), (1,0), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('FONTSIZE', (1,0), (1,0), 10),
    ]))
    Story.append(conclusion_table)
    Story.append(Spacer(1, 2)) 

    # ------------------------------------------------------------------
    # --- VII. PARTICIPATION AUX FRAIS TECHNIQUE -----------------------
    # ------------------------------------------------------------------
    Story.append(Paragraph("<b>VII. PARTICIPATION AUX FRAIS TECHNIQUE</b>", style_heading2))

    fee_result = data.get('fee_result', {})
    fee_check_name = "Contribution aux frais de logistique et d'intervention technique"
    
    # Utilisation des valeurs traitées pour l'affichage (mise à jour pour ne garder que le montant)
    montant_display = fee_result.get('montant_display', 'N/A')
            
    
    # Mise à jour des données du tableau pour ne garder que 2 colonnes
    table_fee_data = [
        [
            Paragraph("<b>Test</b>", style_bold),
            Paragraph("<b>Montant (TTC)</b>", style_bold), # En-tête mis à jour
        ],
        [
            Paragraph(fee_check_name, style_normal),
            Paragraph(montant_display, style_normal),
        ]
    ]
    # Mise à jour des largeurs de colonne pour 2 colonnes
    table_fee_col_widths = [450, 110] 

    table_fee = Table(table_fee_data, colWidths=table_fee_col_widths)
    # Mise à jour du style pour l'alignement de la dernière colonne (la 1)
    style_fee = list(base_table_commands)
    style_fee.append(('ALIGN', (1, 1), (1, -1), 'CENTER')) # La colonne 1 est centrée
    table_fee.setStyle(TableStyle(style_fee))
    Story.append(table_fee)
    Story.append(Spacer(1, 8)) # Laissons cet espace avant le dernier bloc pour aérer le bas de page.
    
    # ------------------------------------------------------------------
    # --- VIII. CLAUSE DE NON-RESPONSABILITÉ ET SIGNATURES ---
    # ------------------------------------------------------------------
    
    # Paragraph 1: Disclaimer/Responsibility (Taille 7)
    resp_text = "Ce test a été réalisé sous la responsabilité du technicien biomédical de l'association AIMA dans le cadre de son activité de réemploi solidaire. L’association AIMA ne peut être tenue pour responsable des éventuels dérèglements des paramètres mesurés et dysfonctionnements du dispositif médical concerné, une fois celui-ci sorti de ses locaux. AIMA décline également toute responsabilité quant aux dommages matériels et/ou corporels qui pourraient survenir lors de la manutention, le transport et l’utilisation ultérieure du dispositif médical."
    
    # Paragraph 2: Humatem Resources (Taille 7)
    humatem_text = "Des fiches-infos matériels, des procédures et des tutoriels vidéo de maintenance sont élaborés et mis gratuitement à disposition par l’ONG Humatem sur son site internet. Une bibliothèque de documentation technique est également disponible en ligne où vous trouverez par exemple des manuels d’utilisation et de maintenance). <a href='https://www.humatem.org/'>www.humatem.org</a>, Rubrique « Centre de ressources »"
    
    # Paragraph 3: AIMA Link (Taille 7)
    aima_link_text = "Pour plus d'information sur l'association AIMA et sur son offre de Réemploi Solidaire : <a href='https://www.assoaima.org/'>https://www.assoaima.org/</a>"

    clause_content = [
        Paragraph(resp_text, style_clause_body),
        Paragraph(humatem_text, style_clause_body),
        Paragraph(aima_link_text, style_clause_body),
    ]

    # B. Bloc de Signatures (Aligné à Droite)
    # Remplacé par un tableau 1 colonne (empilé) pour gagner en largeur.
    sig_data = [
        [Paragraph("Signature du Technicien :", style_normal)], 
        [Spacer(1, 30)],  # Espace pour la signature
        [Paragraph("Validation (Date/Cachet) :", style_normal)],
        [Spacer(1, 30)],  # Espace pour la validation
    ]
    # Largeur réduite à 240 (au lieu de 280)
    sig_table = Table(sig_data, colWidths=[240]) 
    
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'RIGHT'), # Tout à droite
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
    ]))
    
    # C. Tableau principal (Clause à gauche, Signature à droite)
    main_table_data = [
        [clause_content, sig_table]
    ]

    # Largeurs ajustées: 320 (Clause) + 240 (Signatures) = 560 (Total)
    main_table = Table(main_table_data, colWidths=[320, 240]) 
    main_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (0,0), 'LEFT'),      # Clause à gauche
        ('ALIGN', (1,0), (1,0), 'RIGHT'),     # Signatures à droite
        ('VALIGN', (0,0), (-1,-1), 'TOP'),    # Alignement vertical haut
        ('LEFTPADDING', (0,0), (0,0), 5), # Ajouter un padding pour ne pas coller au bord
        ('RIGHTPADDING', (0,0), (-1,0), 5), # Ajouter un padding pour ne pas coller au bord
        ('GRID', (0,0), (-1,-1), 1, colors.black), # CRÉATION DU CADRE (BORDURE NOIRE)
    ]))
    
    Story.append(main_table)
    Story.append(Spacer(1, 20))


    # NOUVEAU: Appel du footer (logos) pour la première page, puis le numéro de page sur toutes les autres
    doc.build(Story, 
              onFirstPage=lambda c, d: (first_page_header_logo(c, d, doc.config_name), first_page_footer_logos(c, d, doc.config_name)), 
              onLaterPages=lambda c, d: later_page_header_logo(c, d, doc.config_name))
    
    buffer.seek(0)
    return buffer


# --------------------------------------------------------------------------------------
# --- FONCTION DE GÉNÉRATION DOCX (CORRIGÉE ET MISE À JOUR POUR LOGOS) -----------------
# --------------------------------------------------------------------------------------

def set_cell_background(cell, color_hex):
    """Définit la couleur d'arrière-plan d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), color_hex[1:].lower()) # ReportLab utilise #RRGGBB, docx RRGGBB
    tcPr.append(shading)

def add_styled_paragraph(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0):
    """Ajoute un paragraphe stylé à une cellule ou au document."""
    if hasattr(cell, 'paragraphs'):
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    else: # Directement au document
        paragraph = cell.add_paragraph()
        
    paragraph.clear()
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(size)
    if bold:
        run.bold = True
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    return paragraph

def add_title(doc, text, level=2):
    """Ajoute un titre H2 ou H3 stylé. Espacement réduit."""
    if level == 2:
        p = doc.add_paragraph()
        p.add_run(text).bold = True
        p.style.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(8) 
        p.paragraph_format.space_after = Pt(1) 
        p.add_run('\n')
    else: # level 3
        p = doc.add_paragraph()
        p.add_run(text).bold = True
        p.style.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1) 

def generate_word_report(data: Dict[str, Any], config: Dict[str, Any]):
    document = Document()
    
    # --- XML HELPER POUR LES BORDS ---
    def set_cell_top_border_xml(cell, border_val='single'):
        """Définit ou retire la bordure supérieure d'une cellule via XML ('nil' pour retirer)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 1. Get or create <w:tcBdr>
        tcBdr = tcPr.find(qn('w:tcBdr'))
        if tcBdr is None:
            tcBdr = OxmlElement('w:tcBdr')
            tcPr.append(tcBdr)
            
        # 2. Create the <w:top> element
        top_bdr = OxmlElement('w:top')
        
        # 3. Set properties
        top_bdr.set(qn('w:val'), border_val)
        if border_val == 'single':
            top_bdr.set(qn('w:sz'), '12')  # 1.5 pt thickness
            top_bdr.set(qn('w:color'), 'auto')
            
        # 4. Remove any existing <w:top> to prevent duplication
        for child in list(tcBdr):
            if child.tag == qn('w:top'):
                tcBdr.remove(child)
                break
                
        # 5. Append the new <w:top>
        tcBdr.append(top_bdr)
    # --------------------------------------------------------------------------

    # Configuration des marges (similaire à A4)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # --- Header (Contient AIMA et le titre - MODIFIÉ POUR LOGO) ---
    header = document.sections[0].header
    header_table = header.add_table(rows=2, cols=2, width=Inches(7.5)) 
    header_table.allow_autofit = False
    header_table.columns[0].width = Inches(3)
    header_table.columns[1].width = Inches(4.5)
    
    # Cellule 1,1
    # NOUVEAU: Insertion du Logo AIMA à la place du texte
    try:
        # Tente d'insérer l'image
        p_logo = header_table.cell(0, 0).paragraphs[0]
        p_logo.clear()
        run = p_logo.add_run()
        run.add_picture(AIMA_LOGO_PATH, width=Inches(0.8)) # Largeur de 0.8 pouce (ajustez)
        
        # Supprime le paragraphe "Solidarité et coopération..." ou l'adapte
        p_coop = header_table.cell(1, 0).paragraphs[0]
        p_coop.clear()
        run_coop = p_coop.add_run("Contrôle Qualité")
        run_coop.font.size = Pt(8)
    except Exception:
        # Fallback au texte si l'image n'est pas trouvée
        add_styled_paragraph(header_table.cell(0, 0), "CONTROL QUALITÉ AIMA", size=10, bold=True, after=0)
        add_styled_paragraph(header_table.cell(1, 0), "Solidarité et coopération biomédicale", size=8, after=0)
    
    # Cellule 1,2
    add_styled_paragraph(header_table.cell(0, 1), "ATTESTATION DE CONTROLE QUALITÉ ET DES PERFORMANCES", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add_styled_paragraph(header_table.cell(1, 1), f"DISPOSITIF MEDICAL (DM): {config.get('NAME', 'N/A')}", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    
    # Ligne de séparation
    document.add_paragraph('').add_run('—'*70).font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Footer (Bas de page) ---
    footer = document.sections[0].footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(7.5))
    footer_table.columns[0].width = Inches(4.5)
    footer_table.columns[1].width = Inches(3.0)
    
    # NOUVEAU: Bloc de Logos en bas à gauche du footer (Cellule 0, 0)
    try:
        p_logos_footer = footer_table.cell(0, 0).paragraphs[0]
        p_logos_footer.clear()
        run_footer = p_logos_footer.add_run()
        run_footer.add_picture(BOTTOM_LOGOS_PATH, width=Inches(1.5)) # Largeur de 1.5 pouce (ajustez)
    except Exception:
         # Fallback au texte si l'image n'est pas trouvée
         add_styled_paragraph(footer_table.cell(0, 0), "Association AIMA - Contacter l'équipe technique.", size=7)


    # Numéro de page en bas à droite (Cellule 0, 1)
    p_page_num = footer_table.cell(0, 1).paragraphs[0]
    p_page_num.clear()
    # Ajout du champ de numéro de page
    p_page_num.add_run("Page ")
    run = p_page_num.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE \\* MERGEFORMAT'
    run._element.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)
    p_page_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # --- I. Identification ---
    add_title(document, "I. Identification de l'Équipement", level=2)
    
    # Data extraction (same as PDF)
    fabricant_raw = data.get('fabricant', 'N/A')
    fabricant_display = str(fabricant_raw)
    if fabricant_display == 'Autre':
        fabricant_display = str(data.get('autre_fabricant', 'Non Spécifié'))
    modele = str(data.get('modele', 'N/A'))
    sn = str(data.get('sn', 'N/A'))
    aima_sn = str(data.get('aima_sn', 'N/A'))
    date_str = str(data.get('date', 'N/A')) 
    lieu = str(data.get('lieu', 'N/A'))
    technicien = str(data.get('technicien', 'N/A'))
    ecme = str(data.get('ecme', 'N/A'))

    id_table = document.add_table(rows=5, cols=4)
    id_table.style = 'Table Grid'
    
    # Fusion et titre
    heading_cell = id_table.cell(0, 0)
    other_cells = [id_table.cell(0, i) for i in range(1, 4)]
    heading_cell.merge(other_cells[-1])
    add_styled_paragraph(heading_cell, "Informations Principales de l'Équipement", size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(heading_cell, '#F2F2F2')

    # Lignes de données
    rows_data = [
        ("Fabriquant:", fabricant_display, "Modèle / Référence:", modele),
        ("N° serie:", sn, "N° série AIMA (GDR):", aima_sn),
        ("Date de controle:", date_str, "Lieu de controle:", lieu),
        ("Nom de l'intervenant technique:", technicien, "ECME utilisé:", ecme),
    ]

    for i, row_d in enumerate(rows_data):
        for j in range(4):
            cell = id_table.cell(i + 1, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if j % 2 == 0: # Labels (bold)
                add_styled_paragraph(cell, f"**{row_d[j]}**", size=8, bold=True, before=1, after=1)
                set_cell_background(cell, '#F2F2F2') 
            else: # Values
                add_styled_paragraph(cell, row_d[j], size=8, before=1, after=1)
                
    # --- II. Vérifications Générales et Visuelles (Simplifié) ---
    add_title(document, "II. Vérifications Générales et Visuelles", level=2)
    add_title(document, "II.A Matériovigilance / Alertes de Sécurité", level=3)

    # Matériovigilance Table
    matvig_status_raw = data.get('materiovigilance')
    matvig_status = '❌ NON' if matvig_status_raw == 'non' or matvig_status_raw is None else '✅ OUI'
    matvig_comment = str(data.get('commentaire_matvig', 'N/A'))
    
    matvig_table = document.add_table(rows=2, cols=3)
    matvig_table.style = 'Table Grid'
    
    headers = ["Contrôle", "Statut", "Commentaire / Observation"]
    for i, h in enumerate(headers):
        add_styled_paragraph(matvig_table.cell(0, i), h, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_background(matvig_table.cell(0, i), '#F2F2F2')

    add_styled_paragraph(matvig_table.cell(1, 0), "L'équipement **n'a pas** fait l'objet d'une matériovigilance / alerte :", size=8)
    add_styled_paragraph(matvig_table.cell(1, 1), matvig_status, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(matvig_table.cell(1, 2), matvig_comment, size=8)


    add_title(document, "II.B Contrôles Visuels / Intégrité Physique", level=3)
    
    visual_checks = config.get("VISUAL_CHECKS", [])
    if visual_checks:
        visuel_table = document.add_table(rows=len(visual_checks) + 1, cols=3)
        visuel_table.style = 'Table Grid'
        
        headers = ["Étape de Contrôle", "OK", "Commentaire / Observation"]
        for i, h in enumerate(headers):
            add_styled_paragraph(visuel_table.cell(0, i), h, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_background(visuel_table.cell(0, i), '#F2F2F2')

        for i, check in enumerate(visual_checks):
            key_check = f'visuel_check_{i}'
            key_comment = f'comment_visuel_{check.replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "").lower()}'
            
            val = data.get(key_check, 'N/A')
            response_text = '✅' if isinstance(val, bool) and val else ('❌' if isinstance(val, bool) and not val else 'N/A')
            comment = str(data.get(key_comment, ''))

            add_styled_paragraph(visuel_table.cell(i + 1, 0), f"• {check}", size=8)
            add_styled_paragraph(visuel_table.cell(i + 1, 1), response_text, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_styled_paragraph(visuel_table.cell(i + 1, 2), comment, size=8)
    else:
         document.add_paragraph("Aucun contrôle visuel spécifique défini.", style='Caption').style.font.size = Pt(8)


    # --- III. Tests de Performance ---
    add_title(document, "III. Tests de Performance et Exactitude", level=2)
    
    perf_checks = config.get("PERFORMANCE_CHECKS", {})
    if perf_checks:
        perf_table = document.add_table(rows=len(perf_checks) + 1, cols=4)
        perf_table.style = 'Table Grid'
        
        headers = ["Test Paramètre", "Statut", "Valeur Injectée/Tolérance", "Valeur Lue/Observation"]
        for i, h in enumerate(headers):
            add_styled_paragraph(perf_table.cell(0, i), h, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_background(perf_table.cell(0, i), '#F2F2F2')
            
        i = 0
        for check_name_key, result in data.get('performance_results', {}).items():
            statut = result.get('statut', 'N/A')
            
            if 'valeur' in result and 'injected' in result:
                lue = result['valeur']
                lue_formatted = f"{lue:.2f}".rstrip('0').rstrip('.') if lue is not None and not np.isnan(lue) else 'N/A'
                seuil_text = f"{result['injected']} {result['unit']} (±{result['tolerance']} {result['unit']})"
                statut_display = '✅ OK' if 'Conforme' in statut else ('❌ NON OK' if 'NON CONFORME' in statut else 'N/A')
                
                obs_text = f"Lue: {lue_formatted} {result.get('unit', '')}. "
                comment = result.get('comment', '')
                obs_text += f"Obs: {comment if comment else 'N/A'}"
                
                add_styled_paragraph(perf_table.cell(i + 1, 0), check_name_key, size=8)
                add_styled_paragraph(perf_table.cell(i + 1, 1), statut_display, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                add_styled_paragraph(perf_table.cell(i + 1, 2), seuil_text, size=8)
                add_styled_paragraph(perf_table.cell(i + 1, 3), obs_text, size=8)
                i += 1
    else:
        document.add_paragraph("Aucun test de performance prédéfini pour cet appareil.", style='Caption').style.font.size = Pt(8)

    # --- IV. Tests Spécifiques ---
    add_title(document, "IV. Tests Spécifiques", level=2)
    
    all_specific_checks = {} 
    for name, expected in config.get('SPECIFIC_CHECKS', {}).items():
        all_specific_checks[name] = expected
    for key in data.get('manual_spec_keys', []):
        all_specific_checks[key] = { 
            'check_name': data.get(f"{key}_name", f"Test Manuel {key.split('_')[-1]}"),
            'expected_result': data.get(f"{key}_expected", "N/A (Saisie Manuelle)")
        }

    if all_specific_checks:
        spec_table = document.add_table(rows=len(all_specific_checks) + 1, cols=4)
        spec_table.style = 'Table Grid'
        
        headers = ["Test Spécifique", "Résultat Attendu", "Statut", "Observation"]
        for i, h in enumerate(headers):
            add_styled_paragraph(spec_table.cell(0, i), h, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_background(spec_table.cell(0, i), '#F2F2F2')

        i = 0
        for check_name_key, expected_result_or_data in all_specific_checks.items():
            if isinstance(expected_result_or_data, str):
                check_name = check_name_key
                expected_result = expected_result_or_data
                result_passed = data.get('specific_results', {}).get(check_name, 'N/A')
                comment_key = check_name 
            else:
                check_name = expected_result_or_data['check_name']
                expected_result = expected_result_or_data['expected_result']
                result_passed = data.get('specific_results', {}).get(check_name_key, 'N/A')
                comment_key = check_name_key 

            response_text = '✅ OK' if isinstance(result_passed, bool) and result_passed else ('❌ NON OK' if isinstance(result_passed, bool) and not result_passed else 'N/A')
            comment = str(data.get(f'comment_spec_{comment_key}', ''))

            add_styled_paragraph(spec_table.cell(i + 1, 0), check_name, size=8)
            add_styled_paragraph(spec_table.cell(i + 1, 1), expected_result, size=8)
            add_styled_paragraph(spec_table.cell(i + 1, 2), response_text, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_styled_paragraph(spec_table.cell(i + 1, 3), comment, size=8)
            i += 1
    else:
        document.add_paragraph("Aucun test spécifique prédéfini pour cet appareil.", style='Caption').style.font.size = Pt(8)


    # --- V. Tests de Sécurité Électrique ---
    add_title(document, "V. Tests de Sécurité Électrique", level=2)
    tse_required = config.get("TSE_REQUIRED", True)

    if not tse_required:
        secu_table = document.add_table(rows=2, cols=3)
        secu_table.style = 'Table Grid'
        headers = ["Test Électrique", "Statut", "Commentaire"]
        for i, h in enumerate(headers):
            add_styled_paragraph(secu_table.cell(0, i), h, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_background(secu_table.cell(0, i), '#F2F2F2')
        add_styled_paragraph(secu_table.cell(1, 0), "Tests de Sécurité Électrique (IEC 62353)", size=8)
        add_styled_paragraph(secu_table.cell(1, 1), "N/A", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_styled_paragraph(secu_table.cell(1, 2), f"Ce dispositif ({config['NAME']}) n'est pas soumis à l'IEC 60601-1 et relève d'une autre norme de sécurité (ex : IEC 61010-1 pour Labo/Stérilisation, ou Classe II/batterie sans terre).", size=8)

    else:
        secu_checks = config.get("SECURITY_CHECKS", {})
        if secu_checks:
            secu_table = document.add_table(rows=len(secu_checks) + 1, cols=4)
            secu_table.style = 'Table Grid'
            
            headers = ["Test Électrique", "Limite", "Mesure", "Statut"]
            for i, h in enumerate(headers):
                add_styled_paragraph(secu_table.cell(0, i), h, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_background(secu_table.cell(0, i), '#F2F2F2')

            i = 0
            for check_name in ['Résistance de terre', 'Courant de fuite patient', 'Courant de fuite châssis']:
                if check_name not in secu_checks: continue
                result = data.get('security_results', {}).get(check_name, {})
                
                statut_display = result.get('statut', 'N/A')
                lim_display = result.get('lim_str', 'N/A')
                mesure_display = result.get('mesure_str', 'N/A')
                
                statut_color = '✅ OK' if 'Conforme' in statut_display else ('❌ NON OK' if 'NON CONFORME' in statut_display else 'N/A')

                add_styled_paragraph(secu_table.cell(i + 1, 0), check_name, size=8)
                add_styled_paragraph(secu_table.cell(i + 1, 1), lim_display, size=8)
                add_styled_paragraph(secu_table.cell(i + 1, 2), mesure_display, size=8)
                add_styled_paragraph(secu_table.cell(i + 1, 3), statut_color, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                i += 1
        else:
            document.add_paragraph("Aucun test de sécurité électrique prédéfini pour cet appareil.", style='Caption').style.font.size = Pt(8)


    # --- VI. CONCLUSION ET VALIDATION ---
    add_title(document, "VI. CONCLUSION ET VALIDATION", level=2)
    
    global_conformity = data.get('global_conformity', 'INCONNU')
    final_comment = data.get('final_comment', 'N/A (Rapport généré avant complétion)')
    
    conc_table = document.add_table(rows=2, cols=2)
    conc_table.style = 'Table Grid'
    conc_table.columns[0].width = Inches(1.5)
    
    add_styled_paragraph(conc_table.cell(0, 0), "Statut Final de l'Équipement:", size=8, bold=True)
    add_styled_paragraph(conc_table.cell(0, 1), global_conformity, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_styled_paragraph(conc_table.cell(1, 0), "Commentaire Final:", size=8, bold=True)
    add_styled_paragraph(conc_table.cell(1, 1), final_comment, size=8)

    # ------------------------------------------------------------------
    # --- VII. PARTICIPATION AUX FRAIS TECHNIQUE -----------------------
    # ------------------------------------------------------------------
    add_title(document, "VII. Participation aux Frais Technique", level=2)

    fee_result = data.get('fee_result', {})
    fee_check_name = "Contribution aux frais de logistique et d'intervention technique"
    
    # Utilisation des valeurs traitées (mise à jour pour l'affichage)
    montant_display = fee_result.get('montant_display', 'N/A')
            
    
    # Mise à jour du tableau pour 2 colonnes
    fee_table = document.add_table(rows=2, cols=2)
    fee_table.style = 'Table Grid'
    fee_table.columns[0].width = Inches(5)
    fee_table.columns[1].width = Inches(2.5)

    headers = ["Test", "Montant (TTC)"] # En-tête mis à jour
    for i, h in enumerate(headers):
        add_styled_paragraph(fee_table.cell(0, i), h, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_background(fee_table.cell(0, i), '#F2F2F2')

    add_styled_paragraph(fee_table.cell(1, 0), fee_check_name, size=8)
    add_styled_paragraph(fee_table.cell(1, 1), montant_display, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ------------------------------------------------------------------
    # --- VIII. CLAUSE DE NON-RESPONSABILITÉ ET SIGNATURES ---
    # ------------------------------------------------------------------
    # Suppression du titre de section
    document.add_paragraph().paragraph_format.space_before = Pt(8)
    
    # Tableau principal 1x2 pour aligner la clause à gauche et les signatures à droite
    main_sig_table = document.add_table(rows=1, cols=2)
    main_sig_table.allow_autofit = False
    main_sig_table.columns[0].width = Inches(4.5)
    main_sig_table.columns[1].width = Inches(3.0) 
    
    # AJOUT DU CADRE: Application du style 'Table Grid'
    main_sig_table.style = 'Table Grid' 
    
    # ---------------------------------------------------
    # A. Colonne de Gauche : Clause de Non-Responsabilité (Taille 7)
    # ---------------------------------------------------
    
    clause_cell = main_sig_table.cell(0, 0)
    
    # Paragraph 1: Disclaimer/Responsibility
    resp_text = "Ce test a été réalisé sous la responsabilité du technicien biomédical de l'association AIMA dans le cadre de son activité de réemploi solidaire. L’association AIMA ne peut être tenue pour responsable des éventuels dérèglements des paramètres mesurés et dysfonctionnements du dispositif médical concerné, une fois celui-ci sorti de ses locaux. AIMA décline également toute responsabilité quant aux dommages matériels et/ou corporels qui pourraient survenir lors de la manutention, le transport et l’utilisation ultérieure du dispositif médical."
    add_styled_paragraph(clause_cell, resp_text, size=7, after=2, before=2)

    # Paragraph 2: Humatem Resources
    humatem_text = "Des fiches-infos matériels, des procédures et des tutoriels vidéo de maintenance sont élaborés et mis gratuitement à disposition par l’ONG Humatem sur son site internet. Une bibliothèque de documentation technique est également disponible en ligne où vous trouverez par exemple des manuels d’utilisation et de maintenance). www.humatem.org, Rubrique « Centre de ressources »"
    add_styled_paragraph(clause_cell, humatem_text, size=7, after=2)
    
    # Paragraph 3: AIMA Link
    aima_link_text = "Pour plus d'information sur l'association AIMA et sur son offre de Réemploi Solidaire : https://www.assoaima.org/"
    add_styled_paragraph(clause_cell, aima_link_text, size=7, after=2)
    
    # ---------------------------------------------------
    # B. Colonne de Droite : Signatures (Aligné à droite)
    # ---------------------------------------------------
    
    sig_cell = main_sig_table.cell(0, 1)
    
    # Création du tableau de signatures imbriqué
    sig_inner_table = sig_cell.add_table(rows=4, cols=1) 
    sig_inner_table.allow_autofit = True 
    sig_inner_table.alignment = WD_TABLE_ALIGNMENT.RIGHT 
    
    # Cellule 0: Titre Signature
    add_styled_paragraph(sig_inner_table.cell(0, 0), "Signature du Technicien :", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, before=2, after=2)

    # Cellule 1: Espace pour Signature
    p_sig = sig_inner_table.cell(1, 0).paragraphs[0]
    p_sig.paragraph_format.space_before = Pt(30) # Espace pour la signature
    p_sig.paragraph_format.space_after = Pt(2)
    set_cell_top_border_xml(sig_inner_table.cell(1, 0), border_val='nil') 

    # Cellule 2: Validation (Date/Cachet)
    add_styled_paragraph(sig_inner_table.cell(2, 0), "Validation (Date/Cachet) :", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, before=2, after=2)

    # Cellule 3: Espace pour Validation
    p_val = sig_inner_table.cell(3, 0).paragraphs[0]
    p_val.paragraph_format.space_before = Pt(30) # Espace pour la validation
    p_val.paragraph_format.space_after = Pt(2)
    set_cell_top_border_xml(sig_inner_table.cell(3, 0), border_val='nil') 

    
    document.add_paragraph().paragraph_format.space_after = Pt(20)


    # Sauvegarde du document en mémoire
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


# --------------------------------------------------------------------------------------
# --- FONCTION PRINCIPALE STREAMLIT (main) (LAISSER INCHANGÉ) ---------------------------
# --------------------------------------------------------------------------------------

def main():
    st.set_page_config(layout="wide", page_title="Contrôle Qualité DM")

    if 'initialized' not in st.session_state:
        st.session_state['initialized'] = True
        st.session_state['device_type'] = DEVICE_TYPES[0]
        st.session_state['materiovigilance'] = 'non'
        st.session_state['commentaire_matvig'] = ''
        st.session_state['manual_spec_count'] = 0
        st.session_state['performance_status'] = 'À RENSEIGNER'
        st.session_state['specific_status'] = 'À RENSEIGNER'
        st.session_state['visual_status'] = 'À RENSEIGNER'
        st.session_state['security_status'] = 'À RENSEIGNER'
        st.session_state['global_conformity'] = 'INCONNU'
        st.session_state['final_comment'] = ''
        st.session_state['performance_results'] = {}
        st.session_state['specific_results'] = {}
        st.session_state['security_results'] = {}
        st.session_state['manual_spec_keys'] = []
        # NOUVEAU : Initialisation de la participation aux frais
        st.session_state['frais_technique'] = ''
        st.session_state['fee_result'] = {} # NOUVEAU : Résultat traité des frais


    st.title("🛠️ Outil de Contrôle Qualité des Dispositifs Médicaux")
    st.markdown("---")

    # --- SÉLECTION DU DISPOSITIF ---
    st.sidebar.header("Configuration")
    st.session_state['device_type'] = st.sidebar.selectbox(
        "Type d'Équipement", DEVICE_TYPES, key='select_device'
    )

    if st.session_state['device_type'] == DEVICE_TYPES[0]:
        st.warning("Veuillez sélectionner le type d'équipement pour commencer le contrôle.")
        return

    # Configuration du dispositif
    if st.session_state['device_type'] == "Autre dispositif":
        st.session_state['device_name'] = st.sidebar.text_input("Nom du Dispositif (Manuel)", "")
        current_config = {"NAME": st.session_state['device_name'], "FABRICANT_LIST": ["Autre"], "VISUAL_CHECKS": [], "PERFORMANCE_CHECKS": {}, "SPECIFIC_CHECKS": {}, "SECURITY_CHECKS": {}, "TSE_REQUIRED": True}
    else:
        current_config = QC_CONFIGS[st.session_state['device_type']]
        st.session_state['device_name'] = current_config['NAME']
    
    st.header(f"Contrôle Qualité : {st.session_state['device_name']}")
    st.markdown("---")


    # ----------------------------------------------------------------------
    # I. IDENTIFICATION
    # ----------------------------------------------------------------------
    with st.expander("I. Identification de l'Équipement", expanded=True):
        col_id1, col_id2, col_id3, col_id4 = st.columns(4)

        col_id1.selectbox(
            "Fabricant", 
            current_config.get("FABRICANT_LIST", ["Autre"]),
            key='fabricant'
        )
        
        if st.session_state.get('fabricant') == 'Autre':
            col_id1.text_input("Spécifiez Fabricant", key='autre_fabricant')
        
        col_id2.text_input("Modèle / Référence", key='modele')
        col_id3.text_input("N° de Série", key='sn')
        col_id4.text_input("N° Identification AIMA (GDR)", key='aima_sn')
        
        col_id1.date_input("Date du contrôle", value=date.today(), key='date')
        col_id2.text_input("Lieu de contrôle", key='lieu')
        col_id3.text_input("Intervenant Technique", key='technicien')
        col_id4.text_input("ECME utilisé (Simulateur / Testeur)", key='ecme')


    # ----------------------------------------------------------------------
    # II. VÉRIFICATIONS GÉNÉRALES ET VISUELLES
    # ----------------------------------------------------------------------
    with st.expander("II. Vérifications Générales et Visuelles", expanded=True):
        st.subheader("II.A Matériovigilance / Alertes de Sécurité")
        col_matvig1, col_matvig2 = st.columns([0.2, 0.8])
        
        col_matvig1.radio(
            "L'équipement a-t-il fait l'objet d'une alerte ou matériovigilance ?", 
            options=['non', 'oui'], 
            index=0 if st.session_state.get('materiovigilance') == 'non' else 1,
            format_func=lambda x: 'NON' if x == 'non' else 'OUI (Alerte confirmée)',
            key='materiovigilance'
        )
        
        st.session_state['commentaire_matvig'] = col_matvig2.text_area(
            "Observation / Référence de l'Alerte :",
            value=st.session_state.get('commentaire_matvig', ''),
            key='commentaire_matvig_input',
            placeholder="Si OUI, spécifiez l'alerte et les actions prises. Si NON, mettez N/A.",
        )
        if st.session_state.get('materiovigilance') == 'oui' and not st.session_state['commentaire_matvig']:
             st.error("Veuillez spécifier l'alerte de matériovigilance et les actions prises.")


        st.subheader("II.B Contrôles Visuels / Intégrité Physique")
        visual_status = "OK"
        visual_checks = current_config.get("VISUAL_CHECKS", [])
        
        for i, check in enumerate(visual_checks):
            key_check = f'visuel_check_{i}'
            key_comment = f'comment_visuel_{check.replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "").lower()}'
            
            col_v1, col_v2, col_v3 = st.columns([0.6, 0.15, 0.25])
            col_v1.markdown(f"**{check}**")
            
            if key_check not in st.session_state:
                st.session_state[key_check] = True

            col_v2.checkbox("OK", value=st.session_state[key_check], key=key_check, label_visibility='collapsed')
            
            st.session_state[key_comment] = col_v3.text_input(
                "Commentaire", 
                key=f'input_{key_comment}', 
                value=st.session_state.get(key_comment, ''), 
                label_visibility='collapsed'
            )

            if not st.session_state[key_check]:
                visual_status = "KO"
            
        if not visual_checks:
            st.info("Aucun contrôle visuel spécifique défini pour cet appareil. Considérez-lo comme OK par défaut.")
            
        st.session_state['visual_status'] = visual_status


    # ----------------------------------------------------------------------
    # III. TESTS DE PERFORMANCE
    # ----------------------------------------------------------------------
    with st.expander("III. Tests de Performance et Exactitude", expanded=True):
        st.subheader("Mesures d'Exactitude des Paramètres")

        performance_status = "OK"
        performance_results = {}
        
        perf_checks = current_config.get("PERFORMANCE_CHECKS", {})
        
        col_p1, col_p2, col_p3, col_p4 = st.columns([0.3, 0.25, 0.25, 0.2])
        col_p1.markdown("**Paramètre**")
        col_p2.markdown("**Mesure (saisie)**")
        col_p3.markdown("**Commentaire**")
        col_p4.markdown("**Statut**")
        
        for check_name, check_data in perf_checks.items():
            key_val = f'perf_val_{check_name}'
            key_comment = f'perf_comment_{check_name}'
            
            col1, col2, col3, col4 = st.columns([0.3, 0.25, 0.25, 0.2])
            
            # Affichage du paramètre et de la tolérance dans la colonne 1 pour une meilleure lisibilité dans l'interface
            col1.markdown(f"**{check_name}** (`{check_data['injected']} {check_data['unit']}` ±{check_data['tolerance']} {check_data['unit']})")
            
            # Saisie de la valeur lue
            val_lue_str = col2.text_input(
                f"Saisie {check_name}", 
                key=key_val,
                value=st.session_state.get(key_val, ""),
                label_visibility='collapsed'
            )
            
            # Saisie du commentaire
            current_comment = col3.text_input( 
                f"Commentaire {check_name}", 
                key=key_comment, 
                value=st.session_state.get(key_comment, ""), 
                placeholder="Commentaire...",
                label_visibility='collapsed'
            )


            try:
                val_lue = float(val_lue_str.replace(',', '.')) if val_lue_str else None

                if val_lue is not None:
                    statut, conforme = valider_performance(
                        check_data["injected"], 
                        val_lue, 
                        check_data["tolerance"], 
                        check_data["unit"]
                    )
                    
                    col4.markdown(statut)
                    if not conforme:
                        performance_status = "KO"
                
                    performance_results[check_name] = {
                        "check_name": check_name,
                        "valeur": val_lue, 
                        "injected": check_data["injected"], 
                        "tolerance": check_data["tolerance"],
                        "unit": check_data["unit"],
                        "statut": statut,
                        "comment": current_comment
                    }
                else:
                    col4.info("À saisir")
                    performance_status = "À RENSEIGNER"
                    performance_results[check_name] = {
                        "check_name": check_name, "valeur": None, "injected": check_data["injected"], 
                        "tolerance": check_data["tolerance"], "unit": check_data["unit"],
                        "statut": "À RENSEIGNER", "comment": current_comment
                    }

            except ValueError:
                col4.error("Saisie invalide")
                performance_status = "KO"
                # Assurez-vous que le résultat est stocké même en cas d'erreur
                performance_results[check_name] = {
                    "check_name": check_name, "valeur": None, "injected": check_data["injected"], 
                    "tolerance": check_data["tolerance"], "unit": check_data["unit"],
                    "statut": "NON CONFORME (Erreur Saisie)", "comment": current_comment
                }
            except Exception:
                 # Gérer les autres exceptions
                 col4.error("Erreur de calcul")
                 performance_status = "KO"
                 performance_results[check_name] = {
                    "check_name": check_name, "valeur": None, "injected": check_data["injected"], 
                    "tolerance": check_data["tolerance"], "unit": check_data["unit"],
                    "statut": "NON CONFORME (Erreur Calcul)", "comment": current_comment
                 }

        st.session_state['performance_status'] = performance_status
        st.session_state['performance_results'] = performance_results


    # ----------------------------------------------------------------------
    # IV. TESTS SPÉCIFIQUES
    # ----------------------------------------------------------------------
    with st.expander("IV. Tests Spécifiques", expanded=True):
        st.subheader("Tests de Fonctionnalités, Alarmes et Pathologies")
        
        specific_status = "OK"
        specific_results = st.session_state.get('specific_results', {})
        
        st.markdown("**Tests Pré-définis**")

        for check_name, expected_result in current_config.get("SPECIFIC_CHECKS", {}).items():
            key_check = check_name 
            key_comment = f'comment_spec_{check_name}'
            
            col_s1, col_s2, col_s3 = st.columns([0.45, 0.15, 0.4])
            col_s1.markdown(f"**{check_name}** *(Attendu: {expected_result})*")
            
            passed = col_s2.checkbox("OK", value=specific_results.get(key_check, True), key=key_check, label_visibility='collapsed')
            specific_results[key_check] = passed

            st.session_state[key_comment] = col_s3.text_input(
                "Observation", 
                key=f'input_{key_comment}', 
                value=st.session_state.get(key_comment, ''), 
                label_visibility='collapsed'
            )

            if not passed:
                specific_status = "KO"
                
        # --- Tests Manuels ---
        st.markdown("---")
        st.markdown("**Tests Manuels Supplémentaires**")
        
        manual_spec_keys = st.session_state.get('manual_spec_keys', [])
        
        if st.button("➕ Ajouter un Test Manuel", key="add_manual_spec"):
            new_key = f'manual_spec_{st.session_state["manual_spec_count"] + 1}'
            manual_spec_keys.append(new_key)
            st.session_state['manual_spec_count'] += 1
            st.session_state['manual_spec_keys'] = manual_spec_keys
            st.experimental_rerun() 

        for manual_key in manual_spec_keys:
            col_m1, col_m2, col_m3, col_m4 = st.columns([0.3, 0.3, 0.15, 0.25])

            col_m1.text_input("Nom du Test", key=f'{manual_key}_name', label_visibility='collapsed')
            col_m2.text_input("Résultat Attendu", key=f'{manual_key}_expected', label_visibility='collapsed')
            
            passed_manual = col_m3.checkbox("OK", value=specific_results.get(manual_key, True), key=manual_key, label_visibility='collapsed')
            specific_results[manual_key] = passed_manual
            
            st.session_state[f'comment_spec_{manual_key}'] = col_m4.text_input(
                "Observation", 
                key=f'input_comment_spec_{manual_key}', 
                value=st.session_state.get(f'comment_spec_{manual_key}', ''), 
                label_visibility='collapsed'
            )

            if not passed_manual:
                specific_status = "KO"

        st.session_state['specific_status'] = specific_status
        st.session_state['specific_results'] = specific_results
        
        
        
# ----------------------------------------------------------------------
# V. TEST DE SÉCURITÉ ÉLECTRIQUE
# ----------------------------------------------------------------------

    with st.expander("V. Test de Sécurité Électrique", expanded=True):
    
        tse_required = current_config.get("TSE_REQUIRED", True) 

        if not tse_required:
            st.info(f"Le **Test de Sécurité Électrique (TSE)** selon IEC 62353 n'est pas applicable à cet appareil ({current_config['NAME']}).")
        
            # NOUVEAU: Ajout du sélecteur de norme
            selected_norme = st.selectbox(
                "Sélectionner la norme de sécurité applicable :",
                ALTERNATIVE_NORMS, # Assurez-vous que cette liste est définie !
                key='secu_norme_alternative'
        )
        
            # Affiche la norme choisie pour confirmation
            st.success(f"L'appareil est considéré conforme à la norme: **{selected_norme}**.")

            st.session_state['security_status'] = 'N/A'
            st.session_state['security_results'] = {}
            # Stocke la norme sélectionnée dans la session state
            st.session_state['secu_val_TSE_required'] = selected_norme 
        
        else:
            st.subheader("Mesures de Sécurité Électrique (selon IEC 62353 ou équivalent)")
        
            # --- NOUVELLES COLONNES ET EN-TÊTES (À AJOUTER/MODIFIER AVANT LA BOUCLE) ---
            # Nous ajoutons une colonne pour le statut principal et utilisons une colonne 
            # séparée pour la saisie optionnelle de la mesure
            col_secu1, col_secu2, col_secu_status, col_secu3, col_secu4 = st.columns([0.3, 0.15, 0.2, 0.15, 0.2])
            col_secu1.markdown("**Test Électrique**")
            col_secu2.markdown("**Limite**")
            col_secu_status.markdown("**Résultat**") # Nouvelle colonne pour le statut (N/A, OK, KO)
            col_secu3.markdown("**Mesure (saisie)**") 
            col_secu4.markdown("**Statut Détail**")
            
            security_status = "OK"
            security_results = {}
            
            for check_name, check_data in current_config.get("SECURITY_CHECKS", {}).items():
                key_val = f'secu_val_{check_name}'
                key_status = f'secu_result_{check_name}' # Clé pour le radio button
                type_test = check_data['type']
            
                # --- NOUVELLES COLONNES DANS LA BOUCLE ---
                col_s1, col_s2, col_s_status, col_s3, col_s4 = st.columns([0.3, 0.15, 0.2, 0.15, 0.2])
                
                col_s1.markdown(check_name)
                col_s2.markdown(f"`{check_data['limit']}`")
            
                # 1. LE BOUTON STANDARD (RADIO) AVEC N/A DANS col_s_status
                selected_status = col_s_status.radio(
                    label="Résultat",
                    options=["Conforme", "Non Conforme", "Non Applicable (N/A)"],
                    key=key_status,
                    index=0, # Par défaut à Conforme
                    horizontal=True,
                    label_visibility='collapsed'
                )
                
                # Initialisation des variables
                val_secu = None
                statut_detail = selected_status
                lim_str = check_data['limit']
                mesure_str = "N/A"
                conforme = True
            
                # 2. LOGIQUE CONDITIONNELLE : Afficher la saisie seulement si "Conforme" ou "Non Conforme"
                if selected_status not in ["Non Applicable (N/A)"]:
                    # Affiche la saisie de mesure
                    val_secu_str = col_s3.text_input(
                        f"Saisie {check_name}", 
                        key=key_val, 
                        value=st.session_state.get(key_val, ""), 
                        label_visibility='collapsed'
                    )
                    
                    try:
                        val_secu = float(val_secu_str.replace(',', '.')) if val_secu_str else None
                        
                        if val_secu is not None:
                            # Calcul de la conformité basée sur la mesure
                            statut_detail, conforme, lim_str, mesure_str = valider_securite(
                                val_secu, 
                                check_data.get('limit', SAFETY_LIMITS_COMMON['fuite_chassis']['limit']), 
                                type_test
                            )
                            
                            if not conforme:
                                security_status = "KO"
                                col_s4.error(statut_detail)
                            else:
                                col_s4.success(statut_detail)
                                
                        else:
                            col_s4.info("Saisie Mesure Requise")
                            security_status = "A RENSEIGNER"
                            statut_detail = "À RENSEIGNER"
            
                    except ValueError:
                        col_s4.error("Saisie invalide")
                        security_status = "KO"
                        statut_detail = "NON CONFORME (Erreur Saisie)"
                else:
                    # Si N/A est sélectionné
                    col_s3.markdown("---") # Visuel vide pour la colonne de saisie
                    col_s4.info("N/A")
                    
                    # Si le statut n'est ni OK ni N/A, on le marque KO
                    if selected_status == "Non Conforme":
                        security_status = "KO"
            
                # 3. MISE À JOUR DU DICTIONNAIRE DE RÉSULTATS
                security_results[check_name] = {
                    "check_name": check_name,
                    "valeur": val_secu,
                    "limite": lim_str,
                    "mesure_str": mesure_str,
                    "statut": statut_detail,
                    "is_na": selected_status == "Non Applicable (N/A)", # Indique si N/A a été choisi
                }
            # --- FIN DE LA BOUCLE ---
            
            st.session_state['security_status'] = security_status
            st.session_state['security_results'] = security_results

# ... le reste du code (qui ne change pas) ...
    # ----------------------------------------------------------------------
    # VI. CONCLUSION ET VALIDATION
    # ----------------------------------------------------------------------
    with st.expander("VI. Conclusion et Validation", expanded=True):
        st.subheader("Statut Global")

        # Logique de Conformité Globale
        statuses = {
            "Visuel": st.session_state['visual_status'],
            "Performance": st.session_state['performance_status'],
            "Spécifique": st.session_state['specific_status'],
            "Sécurité": st.session_state['security_status']
        }
        
        if "KO" in statuses.values():
            global_conformity = "NON CONFORME"
            st.error("🚨 L'équipement est NON CONFORME (KO). Vérifiez les sections en erreur.")
        elif "À RENSEIGNER" in statuses.values():
            global_conformity = "INCOMPLET"
            st.warning("⚠️ Saisie INCOMPLÈTE. Finalisez tous les tests.")
        elif "N/A" in statuses.values() and all(s in ["OK", "N/A"] for s in statuses.values()):
            global_conformity = "APTE À L'UTILISATION CLINIQUE"
            st.success("✅ Équipement APTE À L'UTILISATION CLINIQUE.")
        elif all(s == "OK" for s in statuses.values()):
            global_conformity = "APTE À L'UTILISATION CLINIQUE"
            st.success("✅ Équipement APTE À L'UTILISATION CLINIQUE.")
        else:
            global_conformity = "INCONNU"
            st.info("Statut INCONNU. Vérifiez les saisies.")

        st.session_state['global_conformity'] = global_conformity
        
        col_c1, _ = st.columns([0.3, 0.7])
        col_c1.metric("Statut Final", global_conformity)

        st.session_state['final_comment'] = st.text_area(
            "Commentaire Général de l'Intervention (Obligatoire)",
            value=st.session_state.get('final_comment', ''),
            placeholder="Résumez l'état de l'appareil, les réparations effectuées, et l'autorisation de mise en service."
        )

    # ----------------------------------------------------------------------
    # VII. PARTICIPATION AUX FRAIS TECHNIQUE (MISE À JOUR)
    # ----------------------------------------------------------------------
    with st.expander("VII. Participation aux Frais Technique", expanded=True):
        st.subheader("Contribution Logistique et Intervention")
        
        # En-têtes pour Streamlit (2 colonnes maintenant)
        col_f1, col_f2 = st.columns([0.7, 0.3])
        col_f1.markdown("**Frais Technique**")
        col_f2.markdown("**Montant (TTC)**") # En-tête mis à jour
        
        fee_check_name = "Contribution aux frais de logistique et d'intervention technique"
        fee_key = 'frais_technique'

        col1, col2 = st.columns([0.7, 0.3]) # 2 colonnes pour les données
        col1.markdown(fee_check_name)
        
        fee_raw = col2.text_input(
            "Montant en euros", 
            key=fee_key,
            value=st.session_state.get(fee_key, ""),
            placeholder="ex: 50.00",
            label_visibility='collapsed'
        ).strip()

        fee_result = {}
        
        if not fee_raw:
            fee_result = {
                "montant_saisi": None,
                "statut": "NON SAISI",
                "statut_display": "N/A", # Statut retiré
                "montant_display": "N/A"
            }

        else:
            try:
                # Nettoyer et convertir en float
                fee_float = float(fee_raw.replace(',', '.'))
                
                fee_result = {
                    "montant_saisi": fee_float,
                    "statut": "SAISI ET VALIDE",
                    "statut_display": "N/A", # Statut retiré
                    "montant_display": f"{fee_float:.2f} €"
                }
                
            except ValueError:
                fee_result = {
                    "montant_saisi": None,
                    "statut": "FORMAT INVALIDE",
                    "statut_display": "N/A", # Statut retiré
                    "montant_display": fee_raw + " (Invalide)" # Indiquer l'invalidité
                }
                
        st.session_state['fee_result'] = fee_result


    # ----------------------------------------------------------------------
    # GÉNÉRATION DU RAPPORT
    # ----------------------------------------------------------------------
    st.markdown("---")
    st.subheader("Génération du Rapport")

    if st.session_state['device_type'] != DEVICE_TYPES[0]:
        report_data = {
            key: st.session_state.get(key) 
            for key in st.session_state
        }
        report_data['performance_results'] = st.session_state['performance_results']
        report_data['specific_results'] = st.session_state['specific_results']
        report_data['security_results'] = st.session_state['security_results']
        report_data['global_conformity'] = global_conformity
        report_data['fee_result'] = st.session_state['fee_result'] # Ajout du nouveau résultat structuré
        
        device_name_for_file = st.session_state.get('modele', st.session_state['device_name'])
        # Nettoyage du nom de fichier
        file_name_prefix = device_name_for_file.replace(" ", "_").replace("/", "_").replace(".", "").replace(":", "_")
        report_date = st.session_state['date']

        pdf_buffer = generate_pdf_report(report_data, current_config)
        docx_buffer = generate_word_report(report_data, current_config)

        col_dl1, col_dl2 = st.columns(2)
        
        col_dl1.download_button(
            label="📥 Télécharger Rapport PDF (Recommandé)",
            data=pdf_buffer,
            file_name=f"CERTIFICAT_CQ_{file_name_prefix}_{report_date}.pdf",
            mime="application/pdf"
        )

        col_dl2.download_button(
            label="📥 Télécharger Rapport Word (Modifiable)",
            data=docx_buffer,
            file_name=f"RAPPORT_CQ_{file_name_prefix}_{report_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Veuillez sélectionner un type d'équipement pour permettre la génération du rapport.")


if __name__ == '__main__':
    main()
